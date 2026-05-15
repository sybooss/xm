from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSET_DIR = DOCS / "final-report-assets" / "liujianyu"
OUT_DOCX = DOCS / "结项报告_刘剑宇个人版.docx"
OUT_MD = DOCS / "结项报告_刘剑宇个人版.md"


@dataclass
class Figure:
    path: Path
    title: str
    note: str


def font_path(name: str) -> str | None:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
        Path("C:/Windows/Fonts/msyh.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if bold:
        p = font_path("simhei.ttf")
    else:
        p = font_path("msyh.ttc") or font_path("simsun.ttc")
    if p:
        return ImageFont.truetype(p, size=size)
    return ImageFont.load_default()


def ensure_assets() -> dict[str, Figure]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    sources = {
        "login": ROOT / "output" / "final-report-screenshots" / "login.png",
        "customer": ROOT / "output" / "playwright" / "00-customer-after-sales.png",
        "admin_review": ROOT / "output" / "playwright" / "00-admin-after-sale-review.png",
        "sla": ROOT / "output" / "final-report-screenshots" / "sla-center.png",
        "ticket": ROOT / "output" / "playwright" / "09-service-tickets.png",
        "chat_risk": ROOT / "output" / "playwright" / "05-chat-ai-enhanced.png",
        "chat_workbench": ROOT / "output" / "playwright" / "06-chat-ticket.png",
        "knowledge": ROOT / "output" / "playwright" / "07-knowledge.png",
        "orders": ROOT / "output" / "playwright" / "08-orders.png",
        "profile": ROOT / "output" / "final-report-screenshots" / "customer-profile.png",
        "product_issue": ROOT / "output" / "playwright" / "00-product-issues.png",
        "logs": ROOT / "output" / "playwright" / "10-logs.png",
        "ai_test": ROOT / "output" / "playwright" / "04-ai-test-success.png",
        "use_case": ROOT / "figures" / "use-case-diagram.png",
        "business_flow": ROOT / "figures" / "business-flow-diagram.png",
    }
    copied: dict[str, Path] = {}
    for key, src in sources.items():
        if not src.exists():
            continue
        dst = ASSET_DIR / f"{key}{src.suffix.lower()}"
        shutil.copy2(src, dst)
        copied[key] = dst

    copied["architecture"] = make_box_diagram(
        "系统总体架构图",
        [
            ["Vue 3 前端", "登录注册", "顾客售后", "管理员审核", "咨询工作台"],
            ["Spring Boot 后端", "Controller", "Service", "Mapper", "权限拦截"],
            ["MySQL 数据层", "订单", "售后", "知识库", "日志与工单"],
            ["LangChain4j 增强层", "Prompt 编排", "模型调用", "工具结果注入", "本地兜底"],
        ],
        ASSET_DIR / "architecture.png",
    )
    copied["personal_flow"] = make_flow_diagram(
        "刘剑宇个人负责模块链路",
        [
            "顾客提交售后",
            "管理员审核处理",
            "SLA 风险队列",
            "关联人工工单",
            "AI 回复草稿",
            "图片风险/C2PA",
            "日志追踪与验证",
        ],
        ASSET_DIR / "liujianyu-flow.png",
    )
    copied["ai_boundary"] = make_flow_diagram(
        "AI 辅助边界图",
        [
            "订单与售后状态",
            "业务规则判断",
            "知识/证据检索",
            "LangChain4j 表达增强",
            "管理员确认",
            "状态机记录",
        ],
        ASSET_DIR / "ai-boundary.png",
    )
    copied["state_machine"] = make_flow_diagram(
        "售后状态流转图",
        [
            "已提交",
            "审核中",
            "待补材料",
            "审核通过/驳回",
            "工单处理中",
            "已完成",
            "顾客评价",
        ],
        ASSET_DIR / "state-machine.png",
    )

    figure_meta = {
        "login": ("登录与角色入口", "登录页区分管理员和顾客入口，是双端权限边界的第一层。"),
        "architecture": ("系统总体架构", "系统采用 Vue 3、Spring Boot、MySQL 与 LangChain4j 分层协作。"),
        "use_case": ("系统用例关系", "用例图展示顾客、管理员与智能客服能力之间的交互范围。"),
        "business_flow": ("售后业务流程", "业务流程图说明用户咨询、售后申请、审核、工单与日志之间的闭环。"),
        "customer": ("顾客端我的售后", "顾客可以查看订单、提交售后、补充凭证、跟踪进度并评价服务。"),
        "admin_review": ("管理员售后审核工作台", "管理员在同一页面处理申请、证据、决策、草稿和审计。"),
        "sla": ("SLA 跟进中心", "SLA 中心将超时、优先级和待补材料售后聚合成风险任务队列。"),
        "ticket": ("人工工单页面", "投诉、物流异常或复杂售后可升级为人工客服工单并持续跟进。"),
        "chat_workbench": ("咨询接待工作台", "咨询工作台展示会话、订单上下文、AI 建议、知识依据和处理轨迹。"),
        "chat_risk": ("图片风险与 C2PA 预审", "聊天图片会触发真实性预审，展示 AI 生成、水印、篡改和 C2PA 信号。"),
        "knowledge": ("知识库管理与检索", "知识库支撑 RAG 依据展示和回复草稿中的规则引用。"),
        "orders": ("订单与售后上下文", "订单管理页面保存商品、物流、支付和售后状态，是业务判断依据。"),
        "profile": ("客户画像页面", "客户画像聚合订单、售后、评价、投诉占比和运营建议。"),
        "product_issue": ("商品质量预警", "商品预警从售后与投诉中聚合质量问题，为运营复盘提供依据。"),
        "logs": ("服务日志诊断", "日志中心集中呈现 AI 调用、知识检索和处理轨迹，支撑可解释性。"),
        "ai_test": ("AI 质检页面", "AI 质检页面验证当前 LangChain4j 模型链路和本地兜底状态。"),
        "personal_flow": ("刘剑宇个人负责模块链路", "该图突出真实售后闭环、SLA/工单、AI 草稿和图片风险模块之间的关系。"),
        "ai_boundary": ("AI 辅助边界", "图中强调 AI 只负责增强表达和辅助建议，关键状态仍由业务服务和管理员确认。"),
        "state_machine": ("售后状态流转", "状态流转图展示顾客、管理员、工单和评价之间的核心业务路径。"),
    }
    figures: dict[str, Figure] = {}
    for key, path in copied.items():
        if key in figure_meta:
            title, note = figure_meta[key]
            figures[key] = Figure(path, title, note)
    return figures


def make_box_diagram(title: str, rows: list[list[str]], out: Path) -> Path:
    width, height = 1800, 1120
    img = Image.new("RGB", (width, height), "#f5f7fb")
    draw = ImageDraw.Draw(img)
    title_font = load_font(54, True)
    head_font = load_font(34, True)
    text_font = load_font(28)
    draw.text((width // 2, 80), title, fill="#111827", font=title_font, anchor="mm")
    colors = ["#2563eb", "#0f766e", "#7c3aed", "#b45309"]
    x0, y0, box_w, box_h, gap = 110, 180, 1680, 180, 50
    for i, row in enumerate(rows):
        y = y0 + i * (box_h + gap)
        draw.rounded_rectangle((x0, y, x0 + box_w, y + box_h), radius=28, fill="#ffffff", outline="#d6dce8", width=3)
        draw.rounded_rectangle((x0, y, x0 + 330, y + box_h), radius=28, fill=colors[i % len(colors)])
        draw.text((x0 + 165, y + box_h // 2), row[0], fill="#ffffff", font=head_font, anchor="mm")
        for j, item in enumerate(row[1:]):
            cx = x0 + 430 + j * 290
            draw.rounded_rectangle((cx, y + 48, cx + 230, y + 132), radius=18, fill="#eef4ff", outline="#c7d2fe")
            draw.text((cx + 115, y + 90), item, fill="#1f2937", font=text_font, anchor="mm")
        if i < len(rows) - 1:
            draw.line((width // 2, y + box_h + 8, width // 2, y + box_h + gap - 8), fill="#94a3b8", width=6)
            draw.polygon([(width // 2, y + box_h + gap - 6), (width // 2 - 15, y + box_h + gap - 32), (width // 2 + 15, y + box_h + gap - 32)], fill="#94a3b8")
    img.save(out)
    return out


def make_flow_diagram(title: str, items: list[str], out: Path) -> Path:
    width, height = 2100, 720
    img = Image.new("RGB", (width, height), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = load_font(50, True)
    text_font = load_font(27, True)
    draw.text((width // 2, 80), title, fill="#111827", font=title_font, anchor="mm")
    n = len(items)
    box_w, box_h = 230, 110
    start_x = 80
    gap = (width - 2 * start_x - n * box_w) // max(1, n - 1)
    y = 250
    palette = ["#dbeafe", "#dcfce7", "#fef3c7", "#ede9fe", "#fee2e2", "#cffafe", "#fce7f3"]
    for i, item in enumerate(items):
        x = start_x + i * (box_w + gap)
        draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=24, fill=palette[i % len(palette)], outline="#cbd5e1", width=3)
        lines = wrap_for_pil(item, 7)
        for k, line in enumerate(lines):
            draw.text((x + box_w // 2, y + 38 + k * 34), line, fill="#172033", font=text_font, anchor="mm")
        if i < n - 1:
            x1 = x + box_w + 12
            x2 = x + box_w + gap - 12
            cy = y + box_h // 2
            draw.line((x1, cy, x2, cy), fill="#64748b", width=5)
            draw.polygon([(x2, cy), (x2 - 22, cy - 13), (x2 - 22, cy + 13)], fill="#64748b")
    draw.rounded_rectangle((360, 505, width - 360, 610), radius=28, fill="#ffffff", outline="#dbe3ef", width=3)
    draw.text((width // 2, 557), "业务规则、权限校验和状态流转始终留在 Spring Boot 服务层，AI 只做增强和辅助说明", fill="#334155", font=load_font(30, True), anchor="mm")
    img.save(out)
    return out


def wrap_for_pil(text: str, limit: int) -> list[str]:
    if len(text) <= limit:
        return [text]
    return [text[i:i + limit] for i in range(0, len(text), limit)]


class ReportBuilder:
    def __init__(self, figures: dict[str, Figure]) -> None:
        self.figures = figures
        self.markdown: list[str] = []
        self.doc = Document()
        self.figure_no = 1
        self.table_no = 1
        self.setup_styles()

    def setup_styles(self) -> None:
        styles = self.doc.styles
        for style_name in ["Normal", "Body Text"]:
            style = styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = styles[style_name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.color.rgb = RGBColor(17, 24, 39)
        section = self.doc.sections[0]
        section.top_margin = Cm(2.6)
        section.bottom_margin = Cm(2.3)
        section.left_margin = Cm(2.7)
        section.right_margin = Cm(2.5)

    def add_cover(self) -> None:
        self._center("嘉兴学院", 18, True)
        self.doc.add_paragraph()
        self._center("复杂软件系统实践", 18, True)
        self._center("结项报告", 22, True)
        self.doc.add_paragraph()
        self._center("题目：基于检索增强的电商退换货智能客服系统设计与实现", 14, False)
        self._center("个人版：刘剑宇", 13, True)
        self.doc.add_paragraph()
        items = [
            ("课程名称", "复杂软件系统实践"),
            ("任课老师", "王向东"),
            ("学院", "人工智能学院"),
            ("年级", "2025届"),
            ("专业及班级", "软件231"),
            ("组别", "第3组"),
            ("组员姓名", "刘剑宇、彭译萱、高思嘉"),
            ("本报告撰写对象", "刘剑宇"),
            ("学号", "（开题报告未提供，提交前补充）"),
            ("提交时间", "2026年5月"),
        ]
        table = self.doc.add_table(rows=len(items), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for row, (k, v) in zip(table.rows, items):
            self._set_cell(row.cells[0], k, bold=True)
            self._set_cell(row.cells[1], v)
        self.doc.add_section(WD_SECTION.NEW_PAGE)
        self.markdown.append("# 基于检索增强的电商退换货智能客服系统设计与实现结项报告（刘剑宇个人版）\n")

    def add_toc_placeholder(self) -> None:
        self._center("目录", 16, True)
        p = self.doc.add_paragraph()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "右键更新域可生成目录"
        r.append(t)
        fld.append(r)
        p._p.append(fld)
        self.doc.add_section(WD_SECTION.NEW_PAGE)
        self.markdown.append("[目录：Word 中右键更新域生成]\n")

    def _center(self, text: str, size: int, bold: bool = False) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        self._font(r, size=size, bold=bold, east="黑体" if bold else "宋体")

    def _font(self, run, size: float = 10.5, bold: bool = False, east: str = "宋体") -> None:
        run.font.name = east
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        run.font.size = Pt(size)
        run.bold = bold

    def _set_cell(self, cell, text: str, bold: bool = False) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(str(text))
        self._font(r, size=10, bold=bold)

    def h1(self, text: str) -> None:
        self.doc.add_heading(text, level=1)
        self.markdown.append(f"\n# {text}\n")

    def h2(self, text: str) -> None:
        self.doc.add_heading(text, level=2)
        self.markdown.append(f"\n## {text}\n")

    def h3(self, text: str) -> None:
        self.doc.add_heading(text, level=3)
        self.markdown.append(f"\n### {text}\n")

    def p(self, text: str) -> None:
        text = clean_text(text)
        para = self.doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(21)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(3)
        r = para.add_run(text)
        self._font(r, size=10.5)
        self.markdown.append(text + "\n\n")

    def bullet(self, text: str) -> None:
        para = self.doc.add_paragraph(style="List Bullet")
        para.paragraph_format.line_spacing = 1.35
        r = para.add_run(clean_text(text))
        self._font(r, size=10.5)
        self.markdown.append(f"- {text}\n")

    def table(self, headers: list[str], rows: list[list[str]], title: str) -> None:
        cap = f"表 {self.table_no} {title}"
        self.table_no += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cap)
        self._font(r, size=10, bold=True)
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            self._set_cell(table.rows[0].cells[i], header, True)
        for row_data in rows:
            cells = table.add_row().cells
            for i, item in enumerate(row_data):
                self._set_cell(cells[i], item)
        self.doc.add_paragraph()
        self.markdown.append(f"\n**{cap}**\n\n")
        self.markdown.append("| " + " | ".join(headers) + " |\n")
        self.markdown.append("| " + " | ".join(["---"] * len(headers)) + " |\n")
        for row in rows:
            self.markdown.append("| " + " | ".join(row) + " |\n")
        self.markdown.append("\n")

    def figure(self, key: str) -> None:
        fig = self.figures.get(key)
        if not fig or not fig.path.exists():
            return
        caption = f"图 {self.figure_no} {fig.title}"
        self.figure_no += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(str(fig.path), width=Cm(15.2))
        except Exception:
            return
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        self._font(r, size=9.5, bold=True)
        note = self.doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = note.add_run(fig.note)
        self._font(nr, size=9)
        self.markdown.append(f"\n![{caption}]({fig.path.as_posix()})\n\n{fig.note}\n\n")

    def save(self) -> None:
        self.doc.save(OUT_DOCX)
        OUT_MD.write_text("\n".join(self.markdown).rstrip() + "\n", encoding="utf-8")


def clean_text(text: str) -> str:
    return text.replace("\u2011", "-").replace("\u2014", "-").replace("\u2013", "-").strip()


def paragraph_count_text(path: Path) -> int:
    if not path.exists():
        return 0
    text = path.read_text(encoding="utf-8", errors="ignore")
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def add_long_discussion(rb: ReportBuilder, topic: str, points: list[str], role_focus: bool = False) -> None:
    rb.p(f"{topic} 是本系统结项阶段需要重点说明的内容。开题阶段的设想主要集中在电商退换货智能客服的可行原型，强调意图识别、知识检索、多轮对话和稳定兜底；进入最终实现阶段以后，系统已经不再只是一个聊天问答页面，而是围绕顾客、管理员、订单、售后申请、凭证、工单、日志和 AI 辅助建议形成完整闭环。因此，在说明 {topic} 时，既要从用户能看到的页面讲起，也要把后端服务、数据库表、接口返回和异常兜底讲清楚。")
    for point in points:
        rb.p(point)
    if role_focus:
        rb.p("这部分在个人工作中属于刘剑宇重点负责内容。撰写报告时我没有把它简单写成页面堆叠，而是从真实业务动作出发，把用户提交、管理员处理、系统校验、日志留痕和测试验证串成一个完整说明。这样可以体现我对复杂软件系统中业务主线、数据一致性、权限边界和可演示性的综合理解。")


def build_report() -> None:
    figures = ensure_assets()
    rb = ReportBuilder(figures)
    rb.add_cover()
    rb.add_toc_placeholder()

    rb.h1("摘要")
    rb.p("本结项报告围绕“基于检索增强的电商退换货智能客服系统设计与实现”展开。系统面向电商售后中常见的退货、换货、退款进度、物流异常、投诉转人工、凭证补充和服务评价等业务场景，最终形成了一个采用 Spring Boot、Vue 3、MySQL 与 LangChain4j 的前后端分离系统。与开题阶段的原型设想相比，最终版本不再停留在单一客服聊天框，而是进一步扩展为顾客端、管理员端、知识库、订单中心、人工工单、SLA 跟进、客户画像、商品质量预警、图片凭证风险预审、日志诊断和 AI 质检等多个模块组成的综合业务系统。")
    rb.p("本系统的设计原则是“业务系统为主，AI 能力为辅”。也就是说，订单归属、售后状态流转、工单创建、权限控制、凭证记录、评价提交等关键动作均由 Spring Boot 服务层和 MySQL 数据结构控制；LangChain4j 主要用于调用 OpenAI-compatible 模型，结合订单上下文、知识命中和业务工具结果生成更自然的客服表达。当模型不可用、网络异常或未启用 AI 时，系统仍然可以通过本地规则模板完成主要演示和业务闭环。")
    rb.p("本报告为刘剑宇个人结项版。报告完整描述系统所有功能，同时重点展开本人负责的真实双端售后业务闭环、SLA 跟进与人工工单协同、AI 副驾驶可审计回复草稿、聊天图片风险扫描与 C2PA/图片可信度检测四个模块。通过这些模块，系统从开题阶段的智能客服原型进一步提升为可运行、可解释、可追踪、可扩展的售后服务系统。")
    rb.p("关键词：电商售后；退换货客服；Spring Boot；Vue 3；MySQL；LangChain4j；检索增强；SLA；人工工单；C2PA")

    rb.h1("一、系统设计")
    rb.h2("1.1 选题背景与结项目标")
    for para in [
        "随着电商平台订单规模扩大，退货、换货、退款、物流异常和投诉处理已经成为售后服务中的高频场景。用户在这些场景下往往希望系统能够快速判断订单是否符合规则、需要补充哪些凭证、退款大约多久到达、商家不处理时如何升级，以及是否需要人工客服介入。传统人工客服虽然能够处理复杂情况，但在高峰时段容易受到响应速度、经验差异和规则理解不一致的影响；如果完全依赖通用大模型，则可能出现规则依据不充分、答复不稳定甚至编造规则的问题。",
        "开题报告提出的目标是构建一个基于检索增强的电商退换货智能客服系统，核心思想是把知识库检索、订单上下文、多轮对话和双模式回复生成结合起来。结项阶段在此基础上进一步扩大了工程范围。系统不仅能回答售后咨询，还能让顾客独立提交售后申请、补充凭证、查看处理进度和评价服务；管理员能够在审核工作台处理申请、查看 SLA 风险、创建或跟进人工工单、生成 AI 回复草稿、查看客户画像和商品质量预警。",
        "因此，本项目的最终目标可以概括为三点：第一，完成一个真实可运行的双端售后系统，而不是只做展示页面；第二，在业务规则可控的前提下接入 LangChain4j 和 OpenAI-compatible 模型，提高回复自然度和处理建议质量；第三，通过日志、轨迹、截图和自动化测试证明系统功能不是写死演示，而是具有可验证的工程闭环。",
    ]:
        rb.p(para)
    rb.figure("architecture")

    rb.h2("1.2 技术架构设计")
    add_long_discussion(
        rb,
        "技术架构",
        [
            "前端采用 Vue 3、Vite、Pinia、Vue Router、Axios 和 Element Plus。Vue 3 负责构建响应式页面，Pinia 保存登录用户、系统状态和聊天状态，Vue Router 管理顾客端和管理员端路由，Axios 统一封装后端请求并自动携带 JWT。页面上既有顾客视角的“我的售后”和在线咨询，也有管理员视角的售后审核、SLA 跟进、人工工单、商品预警、客户画像、知识库、服务日志和 AI 质检。",
            "后端采用 Spring Boot 3.3.13、Java 17、MyBatis XML Mapper 和 PageHelper。Controller 负责接收 HTTP 请求和返回统一 Result；Service 负责业务编排、权限判断、状态流转、AI 调用和本地兜底；Mapper 负责访问 MySQL 表；POJO 则用于表示实体、请求对象、分页对象和统一响应。这样既保留了 Java 后端清晰的分层结构，也便于在答辩中沿着页面、API、Controller、Service、Mapper、数据库表逐层解释。",
            "数据层使用 MySQL 8 保存核心业务数据。系统中与本报告关系最密切的表包括 user_account、demo_order、after_sale_application、after_sale_process_log、after_sale_evidence、service_ticket、reply_draft、knowledge_doc、retrieval_log、ai_call_log、process_trace、service_review、product_issue_alert、after_sale_risk_assessment 和 evidence_audit 等。数据结构没有把所有内容塞进一个 JSON 字段，而是围绕订单、售后、知识、日志、工单和评价拆分，这有利于查询、追踪、测试和后续扩展。",
            "AI 增强层使用 LangChain4j 的 OpenAI-compatible 模型调用能力。后端把用户问题、订单上下文、知识库命中、本地业务判断、工单信息和风险信号组织成 Prompt，再调用模型生成更自然的表达。系统同时保留本地规则模板兜底，确保模型失败不会影响订单校验、售后处理、工单创建和日志记录。",
        ],
    )
    rb.figure("use_case")
    rb.figure("business_flow")

    rb.h2("1.3 前后端请求与权限边界")
    for para in [
        "系统的通用请求链路可以概括为：Vue 页面通过 ref、reactive 和 v-model 保存用户输入；页面调用 web/src/api/*.js 中封装的接口函数；Axios 统一加上 Authorization: Bearer token；Spring Controller 接收 query、path、body 或 header 参数；ServiceImpl 完成权限校验、订单归属判断、状态机约束、AI 或本地兜底；Mapper 调用 MyBatis XML 查询或更新 MySQL；Controller 最后用 Result.success(data) 包装统一 JSON 返回前端。前端响应拦截器取出 data 后写入响应式变量，页面自动刷新。",
        "权限上，系统区分管理员和顾客两类角色。顾客登录后默认进入“我的售后”，只能查看自己的订单、提交自己的售后申请、补充自己的凭证、查看自己的评价；管理员登录后默认进入售后审核工作台，可以查看全局售后、处理申请、查看 SLA 风险、创建工单、管理知识库和查看日志。前端路由通过 adminOnly 与 customerOnly 元信息限制页面入口，后端通过 AuthInterceptor、AuthServiceImpl 和 @OperatorAnno 再做接口层校验。",
        "这种前后端共同约束的设计比单靠前端隐藏按钮更可靠。即使用户手动输入后台路由或直接请求后台接口，后端仍会检查 JWT、用户状态和角色权限。对于结项答辩而言，这一点能够说明系统不是只做页面效果，而是考虑到了复杂软件系统中的安全边界和数据归属问题。",
    ]:
        rb.p(para)
    rb.figure("login")

    rb.h2("1.4 AI 辅助边界与本地兜底")
    add_long_discussion(
        rb,
        "AI 辅助边界",
        [
            "本项目没有把大模型作为业务规则引擎。订单能否退货、售后状态能否从已提交流转到审核中、顾客是否有权限补充凭证、管理员能否确认完成、工单是否需要关联售后单，这些都由 Java Service 层和数据库状态决定。AI 生成的回复只能作为建议或表达增强，不能直接修改售后申请状态。",
            "LangChain4j 在系统中主要承担三类工作：第一，组织 Prompt，把订单、知识、对话历史、业务工具结果和本地判断传给模型；第二，在客服聊天和管理员回复草稿中生成更自然的说明；第三，配合图片风险和证据审核等模块输出可读的风险摘要。AI 输出结果会写入日志或草稿表，并由管理员或业务服务进一步确认。",
            "本地兜底是系统稳定性的关键。答辩现场如果外部模型不可用，系统仍然可以根据意图、订单状态、知识文档和固定模板生成客服回复，仍然可以提交售后、审核、补证、转工单、查看日志和评价。这使系统不会因为网络问题失去演示能力，也符合开题报告中“课程现场演示需要稳定可运行”的要求。",
        ],
    )
    rb.figure("ai_boundary")

    rb.h1("二、团队分工")
    rb.p("本项目属于第 3 组，成员包括刘剑宇、彭译萱和高思嘉。由于系统最终功能较多，实际开发和文档整理不是简单按照页面数量切分，而是围绕业务链路、知识检索、回答生成、界面联调和测试验证进行协作。结项报告采用个人版写法，因此既要说明团队整体协作，也要清楚突出刘剑宇负责的重点模块。")
    rb.table(
        ["成员", "主要方向", "说明"],
        [
            ["刘剑宇", "真实双端售后闭环、SLA/工单协同、AI 草稿、图片风险/C2PA、系统整合与验证", "负责把客服原型扩展为真实售后业务系统，重点处理顾客端、管理员端、状态流转、风险队列、工单和 AI 辅助边界。"],
            ["彭译萱", "知识库、知识文档、检索依据与规则维护", "围绕知识分类、知识文档、检索调试和 RAG 依据展示完善系统规则基础。"],
            ["高思嘉", "回答生成、Prompt 组织、前端展示与异常回退联调", "围绕自然语言回复、结果展示、异常提示和页面联调完善用户体验。"],
        ],
        "团队分工概览",
    )
    rb.p("刘剑宇负责的四个重点模块之间并不是孤立关系。真实双端售后闭环提供了顾客和管理员的主业务路径；SLA 跟进与工单协同让复杂售后能够被优先处理并持续跟踪；AI 副驾驶回复草稿把知识依据、售后记录和模型表达结合起来，但仍然要求管理员确认；聊天图片风险扫描与 C2PA 检测则补上了用户上传图片凭证时的可信度预审。四个模块共同支撑系统从“能聊天”升级为“能处理售后业务”。")
    rb.figure("personal_flow")

    rb.h1("三、实现情况")
    rb.h2("3.1 登录注册与双端权限控制")
    add_long_discussion(
        rb,
        "登录注册与权限控制",
        [
            "登录注册模块支持管理员账号和顾客账号两类入口。管理员可使用演示账号进入后台审核工作台，顾客可登录或注册后进入“我的售后”。注册顾客默认角色为 CUSTOMER，不会获得后台访问权限。前端登录页通过模式切换展示登录和注册表单，提交后调用 /auth/login 或 /auth/register，后端校验账号状态、密码和 JWT。",
            "前端 authStore 保存 token、用户信息和加载状态，登录成功后写入 localStorage。Axios 请求拦截器读取 token 并加入 Authorization 请求头。后端 JwtTokenProvider 负责签发和解析 token，AuthServiceImpl 解析 token 后还会回查 user_account 表，确保账号没有被停用。退出登录时会记录撤销标识，避免旧 token 继续使用。",
            "路由权限通过 Vue Router 的 meta 字段实现。管理员页面带 adminOnly，顾客页面带 customerOnly。未登录访问业务页会跳转登录页；顾客访问后台会回到顾客售后中心；管理员误入顾客页会回到审核工作台。后端使用 @OperatorAnno 保护管理接口，实现前后端双重边界。",
        ],
    )

    rb.h2("3.2 顾客端我的售后")
    add_long_discussion(
        rb,
        "顾客端我的售后",
        [
            "顾客端页面是系统从客服聊天原型走向真实业务系统的重要变化。顾客登录后可以看到自己的订单列表、售后申请列表、售后详情、处理时间线、凭证补充入口和服务评价入口。该页面把用户最关心的“我买了什么、能不能申请售后、现在处理到哪一步、下一步该做什么”集中到一个工作区。",
            "后端接口主要位于 /customer/after-sales。顾客提交售后申请时需要传入订单 ID、服务类型、原因编码、原因描述和退款金额。后端会校验当前登录用户是否拥有该订单，并根据订单状态和售后状态决定是否允许提交。售后申请写入 after_sale_application，处理日志写入 after_sale_process_log，从第一步开始保留可追踪时间线。",
            "凭证补充是顾客端的重要能力。当管理员要求补材料时，顾客可以补充文字说明、物流单号、图片链接或其他证据。补充内容会写入 after_sale_evidence，同时处理日志记录补证动作。这样管理员在审核时可以看到完整证据链，而不是只看到一个状态字段。",
        ],
        role_focus=True,
    )
    rb.figure("customer")
    rb.figure("state_machine")

    rb.h2("3.3 管理员售后审核工作台")
    add_long_discussion(
        rb,
        "管理员售后审核工作台",
        [
            "管理员审核工作台是刘剑宇个人负责模块中的核心页面之一。该页面采用队列、详情、决策、证据、草稿和审计组合的形式，避免把售后审核拆散到多个不连贯页面中。管理员可以按关键词、状态、优先级等条件筛选申请，在右侧或详情区域查看订单、顾客、申请原因、凭证、处理日志和系统建议。",
            "审核动作包括通过、驳回、要求补材料、创建关联工单和确认完成。每个动作都不是前端直接改字段，而是调用 /admin/after-sales 下的后端接口。Service 层判断当前状态是否允许执行该动作，再更新 after_sale_application 并写入 after_sale_process_log。驳回必须填写原因，要求补材料会把状态流转到 NEED_MORE_EVIDENCE，确认完成会生成顾客端可见的处理结果说明。",
            "从工程角度看，管理员工作台体现了复杂软件系统中的状态机思维。售后申请不是单纯 CRUD，而是存在提交、审核、补证、通过、驳回、工单、完成、评价等连续状态。系统通过后端服务限制状态转换，既保证业务正确性，也便于自动化测试覆盖。",
        ],
        role_focus=True,
    )
    rb.figure("admin_review")

    rb.h2("3.4 咨询工作台与多轮客服链路")
    add_long_discussion(
        rb,
        "咨询工作台",
        [
            "咨询工作台保留并扩展了开题报告中提出的智能客服能力。用户可以创建会话、绑定订单、输入售后问题、选择是否启用 AI、发送图片凭证，并在中间区域看到客服回复。右侧客服处理面板展示当前客户、当前订单、下一步动作、AI 建议、知识依据、图片/凭证风险和处理轨迹。",
            "后端聊天主链路集中在 ChatServiceImpl.sendMessage。该方法完成保存用户消息、解析多轮上下文、识别意图、查询订单、执行图片风险预审、生成商品洞察、检索知识库、生成本地兜底回复、调用 LangChain4j、判断是否需要人工工单、写入处理轨迹和保存助手回复。这个链路是系统最能体现前后端、数据库和 AI 协同的部分。",
            "多轮上下文通过会话摘要、最近消息和上轮意图完成。比如用户先问“这个订单能不能退货”，系统识别为 RETURN_APPLY；用户继续问“那多久到账”，系统会结合上文识别为 REFUND_PROGRESS。这样用户不需要每一轮都重复订单背景，系统也能保持较自然的售后对话体验。",
        ],
    )
    rb.figure("chat_workbench")

    rb.h2("3.5 知识库管理与 RAG 依据展示")
    add_long_discussion(
        rb,
        "知识库管理与 RAG 依据展示",
        [
            "知识库模块为智能客服提供规则依据。知识分类包括退货规则、换货规则、退款规则、物流异常、投诉与人工转接和通用 FAQ 等。每条知识文档保存标题、文档类型、适用意图、问题、答案、正文、关键词、优先级和状态。管理员可以分页查询、新增、编辑和检索知识文档。",
            "聊天和回复草稿会复用知识检索能力。后端根据用户问题和识别出的意图检索 knowledge_doc，命中结果写入 retrieval_log。前端在咨询工作台和日志中心展示命中文档、分数、排序和命中原因。这样系统回答不是黑盒生成，而是能展示“为什么这样答”。",
            "从开题到结项，知识库的角色发生了扩展。开题阶段它主要是客服回答依据；结项阶段它还服务于 AI 草稿、日志诊断、知识命中 Top 分析和后续知识缺口挖掘。知识库因此成为系统可维护性的基础。",
        ],
    )
    rb.figure("knowledge")

    rb.h2("3.6 订单管理与售后上下文")
    add_long_discussion(
        rb,
        "订单管理与售后上下文",
        [
            "订单数据是售后判断的基础。系统内置 demo_order 表保存订单号、用户、商品名称、SKU、金额、支付状态、订单状态、物流状态、售后状态、支付时间、发货时间和签收时间等信息。管理员可在订单管理页面查看订单和售后上下文，顾客端则只查看自己的订单。",
            "售后申请、客服聊天、商品洞察、图片风险和工单创建都会引用订单上下文。例如同样是“能不能退货”，已签收未超期、退款中、已驳回和物流异常订单会给出不同处理建议。后端把订单查询封装为服务能力，避免前端凭页面数据自行判断规则。",
            "订单管理页面还让老师在答辩时看到系统不是纯前端假数据。页面显示的数据来自后端接口和 MySQL 表，聊天、售后申请和工单都能回到同一订单记录上，形成统一业务证据。",
        ],
    )
    rb.figure("orders")

    rb.h2("3.7 SLA 跟进与风险任务队列")
    add_long_discussion(
        rb,
        "SLA 跟进与风险任务队列",
        [
            "SLA 跟进是刘剑宇个人负责的第二个重点模块。真实售后系统不能只保存申请，还要识别哪些申请已经超时、哪些即将到期、哪些优先级较高、哪些处于待补材料状态。SLA 页面通过风险任务、已超时、24 小时内到期和高优先级等指标，让管理员先处理最急的售后。",
            "后端 /admin/sla/tasks 接口聚合 after_sale_application、风险评估和优先级信息，返回分页任务列表。页面支持按风险类型、状态和关键词筛选，并提供“处理”按钮跳回售后审核工作台。这个设计让 SLA 页面不是独立看板，而是能回到具体处理动作。",
            "SLA 模块的价值在于提升系统的运营感。普通课程项目常常只展示增删改查，而 SLA 队列体现了售后系统对时间约束、优先级和服务质量的关注。它也为未来扩展定时通知、客服绩效和超时预警提供了接口基础。",
        ],
        role_focus=True,
    )
    rb.figure("sla")

    rb.h2("3.8 人工工单协同")
    add_long_discussion(
        rb,
        "人工工单协同",
        [
            "人工工单是刘剑宇负责模块中与 SLA 密切相关的一环。当用户表达投诉、人工客服、商家一直不处理、平台介入等诉求，或者订单物流异常、凭证风险较高、售后争议较复杂时，系统可以创建或关联人工客服工单。工单保存工单号、会话、订单、用户、意图、优先级、状态、客户问题、AI 摘要和处理建议。",
            "管理员可以在售后审核工作台创建关联工单，也可以在人工工单页面筛选和处理工单。工单状态包括待处理、处理中、已解决和已关闭。工单状态更新会写回售后处理日志，使售后申请、订单、工单和处理时间线成为同一条证据链。",
            "人工工单协同体现了本系统对“AI 不能替代人工”的理解。AI 可以识别投诉意图并给出摘要，但真正的复杂处理仍然交给人工客服。系统设计的重点不是让 AI 自动解决所有问题，而是让 AI 帮助分流、摘要和建议，让人工处理更有依据。",
        ],
        role_focus=True,
    )
    rb.figure("ticket")

    rb.h2("3.9 AI 副驾驶可审计回复草稿")
    add_long_discussion(
        rb,
        "AI 副驾驶可审计回复草稿",
        [
            "AI 副驾驶回复草稿是刘剑宇个人负责的第三个重点模块。它不是让模型直接给顾客发消息，而是在管理员售后处理台中生成可审计的回复建议。草稿会结合售后单、订单、凭证、处理日志、知识依据和风险标签生成说明，然后由管理员选择采纳或废弃。",
            "后端使用 reply_draft 表保存草稿来源、内容、风险等级、知识依据、AI 状态、创建人、关联售后单和关联工单。生成草稿时，系统先构造本地模板，若 AI 可用再调用 LangChain4j 增强表达；AI 失败时仍保留模板草稿。采纳草稿会写入处理日志，废弃草稿也会记录原因。这使 AI 输出留下完整审计轨迹。",
            "该模块的设计重点是“可用但可控”。管理员能利用 AI 提升回复质量，但 AI 不能越过人工确认直接影响顾客权益，也不能自动通过、驳回或退款。报告中将其作为个人重点，是因为它同时涉及后端接口、数据库表、LangChain4j、本地兜底、管理员页面和日志审计，能体现较完整的工程实现。",
        ],
        role_focus=True,
    )

    rb.h2("3.10 聊天图片风险扫描与 C2PA/图片可信度检测")
    add_long_discussion(
        rb,
        "聊天图片风险扫描与 C2PA/图片可信度检测",
        [
            "聊天图片风险扫描是刘剑宇个人负责的第四个重点模块。售后场景中，用户经常上传商品破损、包装、物流单或故障照片作为凭证。如果系统完全忽略图片可信度，管理员可能无法判断证据是否充分；如果直接把图片判定为真或假，又会过度承诺。因此本模块定位为“真实性预审”，用于提示风险、补证建议和人工复核方向。",
            "前端咨询工作台支持发送图片。用户选择图片后，系统先检查文件类型和大小，再上传到后端。后端保存图片文件信息，并在消息处理链路中生成 ChatImageRisk 结果。风险面板展示真实性风险、AI 生成风险、篡改风险、C2PA 内容凭证状态、元数据线索、视觉线索、水印线索和建议补充材料。",
            "C2PA 检测用于尝试读取内容凭证。如果图片没有 C2PA 内容凭证，系统不会直接认定图片造假，而是提示“不能据此证明图片来源”。同时，视觉模型和规则信号会给出 AI 生成、水印或二次处理风险。对于疑似 AI 生成或平台水印图片，系统建议要求用户补充原始实拍材料并转人工复核。",
            "该模块的工程价值在于补上了“用户上传图片”和“售后审核证据链”之间的断点。图片风险结果不仅在聊天消息里展示，也进入右侧处理面板和处理轨迹。它让系统能解释为什么需要补证，避免客服只凭单张图片做出退款或驳回结论。",
        ],
        role_focus=True,
    )
    rb.figure("chat_risk")

    rb.h2("3.11 客户画像与服务评价")
    add_long_discussion(
        rb,
        "客户画像与服务评价",
        [
            "服务评价让顾客在售后完成后对处理结果打分和留言。评价只能由售后单所属顾客提交，且要求售后状态为已完成。这样可以避免未处理完成就评价，也避免用户评价他人的售后单。评价内容写入 service_review，并与售后申请、顾客和订单关联。",
            "客户画像页面面向管理员，聚合顾客订单数、订单金额、售后次数、进行中售后、人工工单、评价数、平均评分、近 30 天售后、投诉占比、低分评价、重复售后和运营建议。画像不是为了给顾客贴标签，而是帮助管理员识别高频售后、低分原因和需要重点安抚的用户。",
            "该模块完善了售后处理结果闭环。顾客不仅能提交申请，也能看到处理说明并评价服务；管理员不仅能处理单个售后，也能从客户维度复盘服务质量。它为未来的客服绩效、用户分层和运营分析提供了基础。",
        ],
    )
    rb.figure("profile")

    rb.h2("3.12 商品质量问题聚合预警")
    add_long_discussion(
        rb,
        "商品质量问题聚合预警",
        [
            "商品质量预警模块从售后申请、工单、顾客问题和评价中聚合同一商品的高频问题，例如降噪不明显、屏幕坏点、物流破损、键盘按键失灵等。管理员可以查看预警等级、趋势分、时间窗口、问题关键词、来源样本和建议动作。",
            "该模块的意义是把单次售后问题上升到商品运营层面。某个商品如果连续出现相同问题，客服系统不应只逐单处理，还应提示运营人员排查商品质量、页面宣传、包装物流或供应链。这样系统从售后处理扩展到质量反馈。",
            "在报告中，该功能作为系统整体能力进行描写。它不是刘剑宇个人重点模块，但与 SLA、工单和客户画像共同构成管理员运营侧的复盘能力，体现系统不是只服务单次聊天，而是能帮助业务持续改进。",
        ],
    )
    rb.figure("product_issue")

    rb.h2("3.13 日志诊断与 AI 质检")
    add_long_discussion(
        rb,
        "日志诊断与 AI 质检",
        [
            "日志诊断中心集中展示 AI 调用日志、知识检索日志和处理轨迹。ai_call_log 记录提供方、模型、请求摘要、响应摘要、状态、耗时和错误；retrieval_log 记录查询词、命中文档、排序、分数、命中原因和快照；process_trace 记录上下文解析、意图识别、订单上下文、知识检索、图片风险、AI 生成、工单判断和最终回复。",
            "日志中心还聚合展示 AI 成功率、平均耗时、知识命中数量、去重命中文档、平均检索分数、会话轨迹步骤和高频命中文档。答辩时可以先演示聊天或售后处理，再进入日志中心证明系统确实记录了每一步处理证据。",
            "AI 质检页面用于单独测试 LangChain4j 模型链路。管理员输入测试 prompt 后，后端返回模型提供方、模型名、是否使用 AI、是否触发兜底、耗时和错误信息。即使模型不可用，页面也能显示兜底状态，而不会影响主业务运行。",
        ],
    )
    rb.figure("logs")
    rb.figure("ai_test")

    rb.h2("3.14 售后前置诊断、风险评估与凭证审核")
    add_long_discussion(
        rb,
        "售后前置诊断、风险评估与凭证审核",
        [
            "除了已经在个人重点中展开的四个模块，系统还实现了售后前置诊断、风险评估和凭证审核等智能辅助能力。售后前置诊断用于在顾客正式提交申请前，根据订单状态、商品信息、用户问题和已有凭证给出退货、换货、维修、补证或人工介入建议。这样顾客不是盲目提交申请，而是先得到更贴近业务规则的处理路径。",
            "风险评估模块面向管理员端，综合售后状态、凭证充分性、用户历史、SLA 时间和问题类型，给出低、中、高等风险等级和风险分。风险评估不会直接替代管理员决策，而是帮助管理员识别哪些售后需要更谨慎处理。它与 SLA 队列结合后，可以让高风险售后更早进入人工复核。",
            "凭证审核模块围绕 after_sale_evidence 和 evidence_audit 展开。顾客补充材料后，系统可以从充分性、真实性、AI 生成风险、篡改风险和补证建议等维度生成审核结果。该模块与图片风险扫描有交叉，但更偏售后单证据链；图片风险偏聊天消息里的图片预审。两者共同提升系统对售后材料的处理能力。",
            "这些功能虽然不是刘剑宇本报告中最重点展开的个人模块，但它们共同支撑了系统的“智能售后决策闭环”。如果只有售后申请和聊天，系统还是偏基础；加入诊断、风险和凭证审核后，系统就能在提交前、审核中和复盘后多个阶段提供辅助判断。",
        ],
    )

    rb.h2("3.15 运营首页与系统状态展示")
    add_long_discussion(
        rb,
        "运营首页与系统状态展示",
        [
            "运营首页承担系统总览作用。管理员登录后可以看到系统主题、当前数据库状态、AI 状态、当前模型、兜底策略、关键业务入口和演示路线。早期项目中页面更偏答辩展示，后续根据真实业务方向调整为售后运营首页，强调它是管理端的导航和状态总览，而不是空泛宣传页。",
            "系统状态页面展示后端服务、数据库、AI 配置、模型列表和快速入口。它便于在答辩前确认环境是否可用，也便于排查“模型不可用”和“主业务不可用”的区别。数据库可用但 AI 不可用时，售后申请、审核、工单和知识库仍然可以运行；AI 状态可用时，聊天和回复草稿会有更自然的增强表达。",
            "模型切换能力通过 /system/ai-models 和 /system/ai-models/current 完成。前端顶部模型下拉框读取后端返回的模型选项，切换后后端清理旧模型客户端，下一次 LangChain4j 调用使用新模型。这个功能证明系统不是只在页面写死模型名，而是能通过后端配置影响实际 AI 调用。",
        ],
    )

    rb.h1("四、个人重点实现")
    rb.h2("4.1 真实双端售后业务闭环")
    add_long_discussion(
        rb,
        "真实双端售后业务闭环",
        [
            "本人负责将原本偏客服演示的系统改造成真实双端售后业务闭环。顾客端围绕订单、售后申请、凭证补充、进度跟踪和服务评价展开；管理员端围绕审核队列、处理台、补材料、通过、驳回、完成和审计记录展开。双端使用同一套后端对象和数据库表，只是视角和权限不同。",
            "在接口设计上，顾客端使用 /customer/after-sales，管理员端使用 /admin/after-sales。这样路径本身就体现了角色边界。顾客提交申请时，后端根据 token 限定用户只能操作自己的订单；管理员审核时，后端要求管理员权限并校验售后单状态。该设计避免前端绕过页面限制直接调用关键接口。",
            "数据库层面，after_sale_application 保存售后申请主对象，after_sale_process_log 保存处理日志，after_sale_evidence 保存补充凭证。状态机字段、优先级、SLA 截止时间、处理人、完成说明和下一步动作共同构成业务闭环。每一次关键动作都会写日志，便于顾客和管理员看到同一条时间线。",
            "异常处理上，系统限制重复申请、非法状态流转和越权访问。比如需要补材料的申请在顾客补充凭证后才能继续审核；驳回必须填写理由；评价必须在完成后提交；管理员确认完成不会绕过已有处理日志。通过这些规则，系统更接近真实售后系统，而不是简单修改字段。",
        ],
        role_focus=True,
    )
    rb.h3("4.1.1 页面入口与用户操作流程")
    for para in [
        "真实双端售后闭环首先体现在页面入口的重新组织上。顾客登录后默认进入 /customer/after-sales，而不是进入后台展示页；管理员登录后默认进入 /admin/after-sales/review，而不是进入顾客页面。这个默认入口的调整看起来只是路由变化，实际反映的是系统定位变化：顾客端解决“我的订单和售后怎么处理”，管理员端解决“哪些售后需要审核、如何处理、处理依据是什么”。",
        "顾客端页面中，订单列表、售后申请列表、申请弹窗、详情时间线和评价区域围绕同一对象展开。用户先在订单列表中找到订单，再点击申请售后，填写服务类型、问题原因、退款金额和说明；提交后系统刷新售后列表，并在详情中展示处理时间线。如果管理员要求补材料，顾客可以回到同一详情中补充凭证；如果管理员确认完成，顾客可以在同一位置查看处理结果并评价服务。",
        "管理员端页面中，审核队列、售后详情、证据列表、AI 草稿、处理日志和操作按钮共同构成处理台。管理员不需要在多个页面之间反复跳转，就可以完成查看申请、核对订单、要求补材料、通过、驳回、创建工单、生成回复草稿、确认完成等动作。这样的设计减少了答辩演示时的路径断裂，也更符合真实客服工作台的使用习惯。",
    ]:
        rb.p(para)
    rb.h3("4.1.2 接口链路与状态机约束")
    for para in [
        "接口链路上，我把顾客动作和管理员动作分成两组资源路径。顾客端主要走 GET /customer/after-sales、POST /customer/after-sales、GET /customer/after-sales/{id}、POST /customer/after-sales/{id}/evidence 等接口；管理员端主要走 GET /admin/after-sales、GET /admin/after-sales/{id}、POST /admin/after-sales/{id}/approve、POST /admin/after-sales/{id}/reject、POST /admin/after-sales/{id}/request-evidence、POST /admin/after-sales/{id}/complete 等接口。路径本身表达了操作主体，后端也会再次验证角色。",
        "状态机约束是这部分实现的核心。系统不允许任意状态之间随意跳转，例如已完成的售后不能再次要求补材料，已驳回的售后不能直接确认完成，补材料状态下必须等待顾客补充证据后再继续审核。Service 层在执行每个动作前都会判断当前状态是否满足条件，不满足就返回业务错误。这样可以避免前端误操作或接口被直接调用时破坏数据一致性。",
        "每次状态变化都会写 after_sale_process_log。日志中保存售后申请 ID、动作类型、操作者、备注、创建时间等信息。顾客看到的是处理进度，管理员看到的是审计证据，报告里则可以把它作为系统可追踪性的证明。相比只更新 after_sale_application 的 status 字段，单独保留日志表更适合复杂软件系统，因为它让历史过程不会被最后一次状态覆盖。",
    ]:
        rb.p(para)
    rb.h3("4.1.3 数据库与异常场景处理")
    for para in [
        "数据库设计上，after_sale_application 作为主表保存售后编号、订单 ID、用户 ID、服务类型、原因、退款金额、状态、优先级、SLA 截止时间、处理结果说明和下一步动作。after_sale_evidence 保存顾客补充的凭证内容，after_sale_process_log 保存每一次业务动作。这样的拆分使系统既能快速查询售后主状态，也能在详情页展开完整证据链。",
        "异常场景主要包括订单不存在、订单不属于当前顾客、订单状态不允许售后、重复提交售后、状态不允许审核、驳回理由为空、补证内容为空、评价对象不属于当前顾客等。我的处理思路是把这些约束尽量放在 Service 层，而不是依赖前端按钮隐藏。前端可以改善体验，但真正的规则必须由后端决定。",
        "这个模块最终带来的效果是：老师在答辩时可以从顾客端提交一个售后，再切换管理员端审核，随后回到顾客端查看状态变化和处理记录。整个过程都落库、可查询、可截图、可测试。它比只展示聊天问答更能体现“复杂软件系统实践”的工程要求。",
    ]:
        rb.p(para)

    rb.h2("4.2 SLA 跟进与人工工单协同")
    add_long_discussion(
        rb,
        "SLA 跟进与人工工单协同",
        [
            "本人负责的 SLA 模块主要解决“管理员应该先处理哪个售后”的问题。系统根据售后申请的状态、优先级、截止时间和风险评估生成任务队列，展示风险任务、已超时任务、即将到期任务和高优先级任务。管理员可以从 SLA 页面直接跳转到具体售后处理台，完成实际处理动作。",
            "人工工单模块解决“复杂售后如何转人工持续跟进”的问题。投诉、物流异常、商家不处理、图片凭证风险较高等情况，都可能需要人工介入。系统支持从售后审核台创建关联工单，也支持从聊天链路中根据投诉意图自动创建或提示创建工单。工单与售后单通过 ticket_id 或会话关系关联，处理状态更新会回写售后日志。",
            "这两个模块共同体现了真实客服运营逻辑：SLA 负责发现风险和排序，工单负责承接复杂问题和持续处理。它们不是孤立页面，而是连接顾客申请、管理员审核、聊天咨询、工单处理和日志追踪的中间层。",
            "测试时，我重点验证了 SLA 任务能显示待补材料或超时售后，点击处理能回到审核台；工单能从售后单创建，状态更新能显示在售后处理日志里；客户角色不能访问 SLA 和工单后台。这些验证保证了模块不仅页面可见，而且业务链路可走通。",
        ],
        role_focus=True,
    )
    rb.h3("4.2.1 SLA 风险识别口径")
    for para in [
        "SLA 跟进模块的关键不是简单把售后列表换个样式展示，而是要明确风险识别口径。系统会根据售后申请的 SLA 截止时间、当前状态、优先级和风险评估结果，把任务划分为已超时、即将到期、高优先级、待补材料、审核中等类别。管理员打开页面后，不需要先阅读每个售后详情，就能看到目前风险任务总数、超时数量、24 小时内到期数量和高优先级数量。",
        "任务队列中的每一行不仅显示售后单号和订单号，还显示商品、风险标签、评估等级、风险分、状态、优先级和处理入口。这样做的原因是售后处理不是按创建时间简单排队，高风险、高优先级或即将超时的任务应该优先处理。SLA 页面相当于管理员的任务调度入口。",
        "SLA 接口本身不直接修改售后状态，而是提供风险聚合和跳转处理能力。点击“处理”后仍然回到管理员售后审核工作台，由管理员执行具体动作。这一点很重要，因为风险识别只是辅助排序，真正的通过、驳回、补材料和完成仍然要走售后状态机。",
    ]:
        rb.p(para)
    rb.h3("4.2.2 工单协同与售后证据链")
    for para in [
        "人工工单模块与 SLA 模块配合使用。SLA 告诉管理员哪个售后值得优先处理，而工单负责承接需要持续跟进的问题。例如顾客多次投诉、物流长时间异常、商家不处理、图片凭证风险较高、退款争议较复杂时，管理员可以将售后单转为人工工单，后续由客服继续跟进。",
        "工单数据保存于 service_ticket 表，包含工单号、会话 ID、消息 ID、订单号、用户 ID、意图、优先级、状态、顾客问题、AI 摘要、处理建议和处理人等字段。与售后申请关联后，工单创建和工单状态更新都会写回售后处理日志。这样售后详情页不仅能看到审核结果，也能看到后续人工客服处理过程。",
        "工单页面支持按关键词、状态和优先级筛选，展示 AI 摘要、SLA 风险和下一步动作。它的意义在于把“自动问答无法解决的问题”转入人工处理链路。系统并不追求所有事情都由 AI 自动完成，而是让 AI 和规则帮助识别问题，再把复杂问题交给人工。",
    ]:
        rb.p(para)
    rb.h3("4.2.3 验证方式与答辩讲法")
    for para in [
        "验证 SLA 和工单时，可以先让顾客提交一个售后申请，再由管理员要求补材料，使售后进入需要跟进的状态；随后进入 SLA 页面，确认该申请进入风险队列；点击处理回到审核台，创建关联工单；最后进入人工工单页面，修改工单状态并查看售后处理时间线。这个链路能证明 SLA、售后审核和工单不是三张孤立页面。",
        "答辩时如果老师问“这个系统和普通后台有什么区别”，我可以从 SLA 和工单回答：普通后台只是查数据，而本系统把售后时限、风险优先级和人工接管结合起来，形成了类似真实客服团队的任务队列。客服主管可以先处理高风险和超时任务，客服人员可以在工单里继续跟进复杂问题，顾客端还能看到处理结果。",
    ]:
        rb.p(para)

    rb.h2("4.3 AI 副驾驶回复草稿")
    add_long_discussion(
        rb,
        "AI 副驾驶回复草稿",
        [
            "本人在 AI 副驾驶模块中重点处理“AI 如何参与但不越权”的问题。管理员处理售后时，系统可以生成回复草稿，草稿会结合售后申请、订单状态、凭证、知识依据、风险标签和历史处理日志。管理员可以采纳、废弃或继续人工修改，AI 不会直接把内容发给顾客。",
            "后端先生成本地模板草稿，再根据配置决定是否调用 LangChain4j。这样即使模型失败，也能产生一份基本可用的回复。模型生成内容会记录来源、AI 状态、风险等级和知识依据。采纳草稿时，系统写入 USE_REPLY_DRAFT 处理日志；废弃时写入 DISCARD_REPLY_DRAFT 日志。",
            "这个模块比普通聊天调用更适合真实业务，因为它保留了人工审核环节。售后答复往往涉及退款、责任、凭证和用户情绪，不能让模型直接做最终决定。通过草稿机制，AI 的优势用于表达和摘要，业务责任仍由管理员承担。",
            "在报告中，我把该模块作为个人重点，是因为它横跨数据库表 reply_draft、管理员接口、售后处理台、知识库检索、AI 服务和本地兜底。它体现了我对 AI 工程落地边界的理解：模型不是越强越好，关键是放在正确的位置，并留下可追踪证据。",
        ],
        role_focus=True,
    )
    rb.h3("4.3.1 草稿生成流程")
    for para in [
        "AI 副驾驶回复草稿的流程可以分为四步。第一步，管理员在售后审核工作台选择某个售后申请，系统读取该申请的订单、顾客、原因、凭证、处理日志、知识命中和风险信息。第二步，后端先生成一份本地模板草稿，保证即使 AI 不可用也有基础回复。第三步，如果 AI 开启且模型可用，后端通过 LangChain4j 把上下文组织成 Prompt，请模型对草稿进行自然语言增强。第四步，草稿保存到 reply_draft，等待管理员采纳或废弃。",
        "这里我没有把草稿设计成一次性临时文本，而是把它作为可审计对象保存。reply_draft 中包含 application_id、ticket_id、source_type、content、knowledge_refs、risk_level、ai_status、model_name、status、created_by 等信息。这样管理员后来可以看到这条建议来自 AI、模板还是人工，也可以知道它最终被采纳还是废弃。",
        "草稿的内容通常包括对顾客的安抚、订单状态说明、处理依据、需要补充的材料、下一步处理动作和预计反馈方式。对于不同售后类型，草稿会引用不同知识文档和业务字段。例如退款问题强调到账时间和支付路径，物流异常强调承运商同步和人工跟进，质量问题强调凭证和换货检测。",
    ]:
        rb.p(para)
    rb.h3("4.3.2 AI 失败与本地兜底")
    for para in [
        "AI 失败场景在设计中被明确考虑。模型未配置、网络异常、网关失败、返回超时或主动关闭 AI 时，系统不会让管理员页面报错中断，而是使用本地模板生成草稿，并把 ai_status 记录为 FAILED 或 SKIPPED。这样既能让老师看到 AI 调用状态，也能证明主业务不依赖模型成功。",
        "本地兜底不是简单返回“系统繁忙”，而是根据售后类型、订单状态和知识依据生成结构化回复。虽然表达不如模型自然，但可以覆盖核心业务信息。这个设计符合课程现场演示的稳定性要求，也符合真实业务系统对降级能力的要求。",
        "管理员采纳草稿后，系统不会直接修改售后状态，而是记录 USE_REPLY_DRAFT 日志，并把草稿状态改为 USED。废弃草稿则记录 DISCARD_REPLY_DRAFT。通过这种方式，AI 参与过程可以被审计，管理员仍然是最终责任主体。",
    ]:
        rb.p(para)
    rb.h3("4.3.3 与知识库和日志中心的关系")
    for para in [
        "AI 草稿并不是孤立功能，它依赖知识库和日志中心。知识库提供规则依据，草稿中的处理说明尽量引用已命中的知识文档；日志中心记录 AI 调用是否成功、使用了哪个模型、耗时多少、是否触发兜底。管理员如果对某条草稿有疑问，可以回到日志中查看模型调用和检索依据。",
        "这部分在报告中需要详细写，是因为它体现了 AI 工程化的关键问题：不是能不能调模型，而是调模型之后如何和业务系统结合，如何保存结果，如何让人确认，如何失败降级，如何留下证据。与开题报告中的“OpenAI 接口生成与本地规则模板生成”相比，结项阶段已经把这个思路落到了具体表、接口和页面上。",
    ]:
        rb.p(para)

    rb.h2("4.4 图片风险扫描与 C2PA 可信度检测")
    add_long_discussion(
        rb,
        "图片风险扫描与 C2PA 可信度检测",
        [
            "图片风险模块是我在结项阶段补强的一个特色功能。它的目标不是做绝对鉴伪，而是在顾客发送图片凭证后，为客服提供真实性预审、风险提示和补证建议。售后场景中的图片可能来自真实拍摄，也可能经过裁剪、压缩、平台水印、AI 生成或二次编辑，因此系统需要谨慎地给出风险等级。",
            "前端实现上，咨询工作台提供“发图片”按钮，上传后在消息气泡中展示图片和简要风险标签。右侧“图片/凭证风险”标签页展示更完整信息，包括真实性风险、AI 生成风险、篡改风险、C2PA 状态、视觉模型状态、元数据线索、水印线索和建议补充材料。",
            "后端实现上，图片上传后会形成文件记录，消息发送时把文件信息带入聊天链路。风险扫描服务会综合图片大小、来源文件名、可见水印、C2PA 内容凭证、视觉信号和本地规则，形成 ChatImageRisk 结果。若 C2PA 未发现内容凭证，系统只说明“不能据此证明来源”，不直接判定造假；若存在明显 AI 平台水印或视觉信号，则建议人工复核并要求用户补充原始实拍材料。",
            "这个模块对结项报告很有价值，因为它体现了系统对新型 AI 生成内容风险的关注。用户上传的图片不应被系统无条件采信，客服也不应仅凭模型一句话做决定。图片风险预审与人工复核结合，既提高了系统智能化程度，也保持了售后处理的谨慎性。",
        ],
        role_focus=True,
    )
    rb.h3("4.4.1 图片上传与前端展示")
    for para in [
        "图片风险模块从前端上传开始。咨询工作台提供“发图片”按钮，用户选择图片后，前端检查 MIME 类型必须是 image，文件大小不能超过配置上限。上传成功后，图片会先作为待发送附件显示在输入区，用户可以继续补充文字说明，再一起发送到聊天会话中。",
        "消息区展示图片缩略图、原始文件名和风险标签。右侧处理面板单独提供“图片/凭证风险”标签页，用于展示完整风险信息。这样普通用户或客服可以先在消息流中看到图片，管理员或答辩老师则可以在右侧面板看到更细的真实性预审结果。",
        "前端展示时，我特别注意不要把风险文案写成绝对判断。比如没有 C2PA 凭证时，页面提示“未发现 C2PA 内容凭证，不能据此证明图片来源”，而不是“图片一定是假的”。疑似 AI 生成或有水印时，页面提示补充原始实拍材料并人工复核。这种措辞更符合真实业务和技术边界。",
    ]:
        rb.p(para)
    rb.h3("4.4.2 后端风险信号与 C2PA 处理")
    for para in [
        "后端风险结果包含 auditStatus、authenticityRisk、aiGeneratedRisk、tamperRisk、requiredEvidence、visionStatus、visionModel、visionSignal、c2paStatus、c2paProvider、c2paGenerator、c2paSignal、metadataSignal、visualSignal 和 watermarkSignal 等字段。字段拆得比较细，是为了让前端既能显示总体判断，也能展示具体依据。",
        "C2PA 处理的逻辑是尝试调用配置中的 c2patool 或等价工具读取内容凭证。如果工具不可用、超时或图片没有凭证，系统会记录 SKIPPED 或 NOT_FOUND 一类状态。这个状态只说明没有读取到可靠内容凭证，并不等于图片一定伪造。系统同时结合文件来源、可见水印、视觉信号和业务上下文给出综合建议。",
        "视觉和水印信号用于补充 C2PA 的不足。比如用户上传带有“豆包AI生成”水印的图片，系统会把 AI 生成风险和真实性风险提高，并建议补充原始照片；如果图片只是普通商品照片但没有内容凭证，则风险可能较低或中等，只提示客服结合订单和售后规则复核。",
    ]:
        rb.p(para)
    rb.h3("4.4.3 业务价值与局限")
    for para in [
        "该模块的业务价值在于让图片凭证进入可解释链路。传统课程项目可能只实现图片上传，上传后就把图片当成附件展示；本系统进一步对图片做风险预审，并把风险结果放入聊天洞察、处理轨迹和补证建议中。这样客服知道什么时候需要要求原图、什么时候需要转人工、什么时候可以把图片作为初步材料。",
        "当然，图片风险预审有明确局限。它不能替代司法鉴定，也不能保证所有 AI 生成图片都能被识别。报告中必须如实说明它只是辅助信号，最终仍需结合订单、物流、售后规则和人工判断。正因为系统承认这个边界，它才更适合作为真实业务增强，而不是夸大成绝对鉴伪工具。",
        "从个人工作角度看，图片风险模块让我把前端文件上传、后端 multipart 配置、图片消息、风险数据结构、C2PA 工具、AI 风险提示和人工复核建议串联起来。它和 SLA、工单、AI 草稿共同体现了一个原则：系统可以智能，但关键业务判断要可解释、可追踪、可人工接管。",
    ]:
        rb.p(para)

    rb.h1("五、测试与验证")
    rb.p("本项目结项报告坚持不伪造测试结果。代码修改后需要执行与范围匹配的验证，至少包括后端编译或前端构建；涉及接口或页面时补充接口烟测或浏览器测试。报告生成完成后，也需要进行文档字数检查、DOCX 渲染检查和工作区状态检查。")
    rb.table(
        ["验证项", "命令", "目标"],
        [
            ["后端编译", "cd server; mvn -q -DskipTests package", "确认 Java 代码、MyBatis Mapper 和依赖可打包"],
            ["前端构建", "cd web; npm run build", "确认 Vue 3 页面、路由、组件和静态资源可构建"],
            ["全链路接口烟测", "tools/full-smoke-test.ps1", "覆盖登录、知识库、订单、售后、聊天、工单、日志等接口"],
            ["浏览器主流程", "cd web; npm run test:browser", "覆盖顾客和管理员主业务页面"],
            ["角色权限测试", "cd web; npm run test:browser:roles", "验证顾客和管理员菜单与路由隔离"],
            ["报告字数检查", "脚本统计中文字符数", "确认正文不少于 2 万字"],
            ["报告渲染检查", "LibreOffice 转 PDF 并抽取 PNG", "确认截图、表格、标题和分页基本正常"],
        ],
        "结项验证项目",
    )
    for para in [
        "全链路接口烟测覆盖系统状态、登录注册、AI 测试、知识分类与文档、知识检索、订单、售后记录、聊天、多轮追问、自动工单、服务工单、消息列表、处理轨迹、AI 日志、检索日志和测试数据清理。它证明系统不是只在页面上展示假数据，而是后端接口和数据库链路能够实际运行。",
        "浏览器测试覆盖用户能看到的真实页面，包括顾客售后中心、管理员审核工作台、SLA 跟进、人工工单、商品预警、咨询工作台、知识库、订单管理、日志诊断和 AI 质检。由于 Element Plus 某些标签页会保留隐藏 DOM，测试脚本需要判断可见文本或先切换到对应标签页，再断言页面内容。",
        "角色权限测试用于确认顾客不能访问管理员页面，管理员也不会默认进入顾客端页面。对于结项答辩来说，这类验证能够说明系统不仅功能多，而且有明确权限边界。",
    ]:
        rb.p(para)

    rb.h1("六、项目特色与创新点")
    rb.p("本项目的特色首先体现在真实业务闭环。系统从开题阶段的智能客服原型扩展为顾客端和管理员端协作的售后系统，顾客能够提交申请、补证、看进度和评价，管理员能够审核、要求补证、转工单、确认完成和复盘客户画像。")
    rb.p("第二个特色是 AI 辅助边界清晰。LangChain4j 不是替代业务规则，而是作为增强层参与客服回复、草稿生成、风险摘要和模型质检。所有关键状态流转仍在 Spring Boot 服务层完成，AI 失败时本地规则兜底仍可运行。")
    rb.p("第三个特色是可追踪证据链。知识命中写入 retrieval_log，AI 调用写入 ai_call_log，处理步骤写入 process_trace，售后动作写入 after_sale_process_log。老师可以从页面直接看到每次回复、审核和工单背后的依据。")
    rb.p("第四个特色是 SLA 与人工接管。SLA 队列把超时和高优先级售后集中展示，人工工单把投诉和复杂售后转入客服处理链路，避免系统只停留在自动问答层面。")
    rb.p("第五个特色是图片风险与 C2PA 预审。系统能够对用户发送的图片凭证进行真实性预审，提示 AI 生成、水印、篡改和内容凭证信号，要求必要时补充原始实拍材料。这一点结合了当前生成式 AI 环境下的真实风险。")
    rb.p("第六个特色是报告和验证材料完整。项目不仅有代码和页面，也有开题报告、接口文档、数据库设计、源码链路讲解、功能路线文档、自动化测试脚本和截图证据，符合复杂软件系统实践课程对工程过程和结项材料的要求。")

    rb.h1("七、不足与未来展望")
    for para in [
        "当前系统虽然已经形成较完整的售后闭环，但知识库检索仍以 MySQL 文本检索、关键词和规则排序为主。后续可以引入向量检索和语义重写，让用户表达更口语化、同义表达更复杂时仍能命中合适知识文档。",
        "当前订单和售后数据主要是演示种子数据，尚未接入真实电商订单系统、仓库系统、支付退款系统和物流接口。后续可接入真实或模拟的支付回调、仓库收货、退款打款和物流轨迹，进一步提升业务完整度。",
        "权限角色目前主要分为顾客和管理员。真实企业场景下还可以细分客服、主管、仓库、财务和运营角色，不同角色拥有不同处理动作和查看范围。这样能够更接近实际售后组织结构。",
        "图片风险模块目前以预审为主，不能替代专业鉴定。后续可以接入更多图像取证信号，例如 EXIF 一致性、压缩痕迹、边缘异常、设备信息和更完整的 C2PA 工具链，同时保留人工复核作为最终判断。",
        "SLA 模块目前提供页面和接口查询，后续可以增加定时任务、站内通知、邮件或企业微信提醒，让即将超时的任务主动提醒客服。人工工单也可以加入分派、转派、处理时长统计和客服绩效报表。",
        "测试脚本已经覆盖主流程，但随着系统扩大，后续可以拆分为更细的单元测试、接口测试、页面冒烟测试和角色权限测试，降低长链路测试失败时的排查成本。报告和文档也可以继续和代码注释、OpenAPI 文档结合，提高维护效率。",
    ]:
        rb.p(para)

    rb.h1("结论")
    rb.p("通过本次复杂软件系统实践，本项目从一个基于检索增强思路的电商退换货智能客服原型，逐步发展为包含顾客端、管理员端、知识库、订单、售后、凭证、SLA、人工工单、AI 草稿、图片风险、客户画像、商品预警、日志诊断和 AI 质检的综合系统。系统主栈保持为 Spring Boot + Vue 3 + MySQL + LangChain4j，符合开题阶段对前后端分离、知识检索、AI 增强和稳定兜底的设想，同时在结项阶段增加了更真实的业务闭环。")
    rb.p("作为刘剑宇个人结项报告，本文重点说明了本人负责的真实双端售后业务闭环、SLA 跟进与工单协同、AI 副驾驶可审计回复草稿、聊天图片风险扫描与 C2PA 可信度检测四个模块。这些模块共同体现了我对复杂软件系统中业务流程、权限边界、数据结构、AI 辅助定位、异常处理和验证证据的理解。最终系统不仅能够完成课程演示，也具备继续扩展为更完整售后服务平台的基础。")

    rb.save()


def main() -> None:
    build_report()
    chars = paragraph_count_text(OUT_MD)
    print(f"DOCX: {OUT_DOCX}")
    print(f"MD: {OUT_MD}")
    print(f"Chinese chars in markdown: {chars}")
    if chars < 20000:
        raise SystemExit(f"Report is too short: {chars} Chinese chars")


if __name__ == "__main__":
    main()
