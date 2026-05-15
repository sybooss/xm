from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSET_DIR = DOCS / "final-report-assets" / "liujianyu"
OUT_DOCX = DOCS / "结项报告_刘剑宇个人版.docx"
OUT_MD = DOCS / "结项报告_刘剑宇个人版.md"
REMOVED_GENERATED_FIGURES = {
    "architecture.png",
    "liujianyu-flow.png",
    "ai-boundary.png",
    "state-machine.png",
}
REAL_FIGURES = [
    ("login.png", "登录与角色入口", "登录页区分管理员和顾客入口，是双端权限边界的第一层。"),
    ("use_case.png", "系统用例关系", "用例图展示顾客、管理员与智能客服能力之间的交互范围。"),
    ("business_flow.png", "售后业务流程", "业务流程图说明用户咨询、售后申请、审核、工单与日志之间的闭环。"),
    ("customer.png", "顾客端我的售后", "顾客可以查看订单、提交售后、补充凭证、跟踪进度并评价服务。"),
    ("admin_review.png", "管理员售后审核工作台", "管理员在同一页面处理申请、证据、决策、草稿和审计。"),
    ("sla.png", "SLA 跟进中心", "SLA 中心将超时、优先级和待补材料售后聚合成风险任务队列。"),
    ("ticket.png", "人工工单页面", "投诉、物流异常或复杂售后可升级为人工客服工单并持续跟进。"),
    ("chat_workbench.png", "咨询接待工作台", "咨询工作台展示会话、订单上下文、AI 建议、知识依据和处理轨迹。"),
    ("chat_risk.png", "图片风险与 C2PA 预审", "聊天图片会触发真实性预审，展示 AI 生成、水印、篡改和 C2PA 信号。"),
    ("knowledge.png", "知识库管理与检索", "知识库支撑 RAG 依据展示和回复草稿中的规则引用。"),
    ("orders.png", "订单与售后上下文", "订单管理页面保存商品、物流、支付和售后状态，是业务判断依据。"),
    ("profile.png", "客户画像页面", "客户画像聚合订单、售后、评价、投诉占比和运营建议。"),
    ("product_issue.png", "商品质量预警", "商品预警从售后与投诉中聚合质量问题，为运营复盘提供依据。"),
    ("logs.png", "服务日志诊断", "日志中心集中呈现 AI 调用、知识检索和处理轨迹，支撑可解释性。"),
    ("ai_test.png", "AI 质检页面", "AI 质检页面验证当前 LangChain4j 模型链路和本地兜底状态。"),
]


def recover_markdown_from_head() -> str:
    paths = subprocess.check_output(
        ["git", "-c", "core.quotepath=false", "ls-tree", "-r", "--name-only", "HEAD"],
        encoding="utf-8",
    ).splitlines()
    needle = "刘剑宇"
    md_path = next(p for p in paths if needle in p and p.endswith(".md"))
    return subprocess.check_output(["git", "show", "HEAD:" + md_path]).decode("utf-8")


def clean_markdown(text: str) -> str:
    text = re.sub(r"!\[[^\]]+\]\([^)]+\)\n\n[^\n]*\n?", "", text)
    lines = []
    for line in text.splitlines():
        if line.startswith("[目录："):
            continue
        lines.append(line)
    text = "\n".join(lines)

    top_matches = list(re.finditer(r"(?m)^# (.+)$", text))
    sections: dict[str, str] = {}
    title = ""
    for idx, match in enumerate(top_matches):
        heading = match.group(1).strip()
        start = match.end()
        end = top_matches[idx + 1].start() if idx + 1 < len(top_matches) else len(text)
        body = text[start:end].strip()
        if idx == 0 and "结项报告" in heading:
            title = heading
            continue
        sections[heading] = body

    def section_body(prefix: str) -> str:
        for heading, body in sections.items():
            if heading.startswith(prefix):
                return body
        return ""

    intro = section_body("摘要").replace("本结项报告围绕", "本课题围绕").replace("关键词：", "主要关键词：")
    team = section_body("二、团队分工")
    design = section_body("一、系统设计")
    implementation_parts = [
        section_body("三、实现情况"),
        section_body("四、个人重点实现"),
        "## 4.5 测试与验证\n\n" + section_body("五、测试与验证"),
        "## 4.6 项目特色与创新点\n\n" + section_body("六、项目特色与创新点"),
        appendix,
    ]
    future = "\n\n".join(part for part in [section_body("七、不足与未来展望"), section_body("结论")] if part)

    text = "\n\n".join(
        part
        for part in [
            f"# 一、课题介绍\n\n{intro}",
            f"# 二、团队分工\n\n{team}",
            f"# 三、系统设计\n\n{design}",
            "# 四、实现情况\n\n" + "\n\n".join(part for part in implementation_parts if part),
            f"# 五、未来与展望\n\n{future}",
        ]
        if part.strip()
    )

    text = text.replace("\n# 五、测试与验证", "\n## 4.5 测试与验证")
    text = text.replace("\n# 六、项目特色与创新点", "\n## 4.6 项目特色与创新点")
    text = text.replace("\n# 四、个人重点实现", "\n## 4.9 个人重点实现")
    text = re.sub(r"\n# 结论\n", "\n## 5.7 总结\n", text)
    return compact_markdown_tables(text.strip()) + "\n"


def compact_markdown_tables(text: str) -> str:
    lines = text.splitlines()
    out: list[str] = []
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith("|"):
            table_lines = []
            while i < len(lines):
                current = lines[i].strip()
                if current.startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                    continue
                if current == "" and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                    i += 1
                    continue
                break
            out.append("\n".join(table_lines))
            continue
        out.append(lines[i])
        i += 1
    return "\n".join(out)


appendix = """

## 4.14 个人负责模块补充说明

为了让个人结项报告与实际分工一致，本报告特别强调刘剑宇负责的四个模块不是相互孤立的页面，而是围绕售后服务闭环形成连续链路。顾客端售后中心解决“用户如何发起和跟踪问题”，管理员审核工作台解决“平台如何判断和处理问题”，SLA 与人工工单解决“风险任务如何被优先处理并转入人工接管”，AI 副驾驶草稿解决“模型如何在不越权的前提下辅助表达”，图片风险与 C2PA 检测解决“用户上传凭证如何进行可信度预审”。这几部分组合起来，才能体现系统从普通聊天工具升级为复杂软件系统。

在实现这些模块时，我最关注的是业务动作之间的数据一致性。比如顾客提交售后后，管理员端必须能看到同一条申请；管理员要求补材料后，顾客端必须显示可补证入口；补证完成后，审核台需要能继续处理；创建人工工单后，工单页面和售后详情都要能看到关联关系；AI 生成草稿后，管理员采纳或废弃都必须留下日志。只有这些状态和日志互相对应，系统才不是静态展示。

我也特别注意 AI 能力的边界。AI 可以辅助生成回复、总结风险、润色表达和解释知识依据，但不能直接批准退款、驳回申请或关闭工单。所有关键状态仍然由 Spring Boot 服务层根据规则和管理员动作完成，数据库记录每一步结果。这样设计的原因是售后业务涉及用户权益和平台责任，不能把最终决策交给不可完全解释的模型。

图片风险模块同样遵循这个原则。C2PA 或图片风险扫描只能提供辅助信号，不能替代人工审核。系统提示“未发现内容凭证”时，不等于证明图片伪造；提示“疑似 AI 生成”时，也需要结合订单、物流、商品和用户补充材料综合判断。报告中这样写，是为了避免把技术能力夸大成绝对鉴定能力。

本次报告修订还调整了截图和图示策略。上一版中由脚本绘制的彩色架构图和流程图虽然能表达结构，但与课程 Word 模板风格不一致，也容易让报告显得像通用宣传材料。因此修订版删除这些自绘图，改用真实系统截图、已有项目图和正文表格说明。若后续确实需要 AI 生成或编辑图片，将严格按照 AGENTS.md 规定，只使用 gpt-image-2。
"""


class TemplateDoc:
    def __init__(self) -> None:
        self.doc = Document()
        self.figure_no = 1
        self.setup()

    def setup(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.6)
        for style_name in ["Normal", "Body Text"]:
            style = self.doc.styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(12)
        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = self.doc.styles[style_name]
            style.font.name = "楷体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
            style.font.color.rgb = RGBColor(0, 0, 0)

    def font(self, run, size: float = 12, bold: bool = False, east: str = "宋体") -> None:
        run.font.name = east
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        run.font.size = Pt(size)
        run.bold = bold

    def add_cover(self) -> None:
        for _ in range(1):
            self.doc.add_paragraph()
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        set_cell_width(cell, 4.8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        p = cell.paragraphs[0]
        r = p.add_run("成绩：\n")
        self.font(r, size=16)
        for _ in range(1):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("嘉兴学院")
        self.font(r, size=42)
        for _ in range(2):
            self.doc.add_paragraph()
        for line in [
            "课程名称：复杂软件系统实践",
            "任课老师：王向东",
            "题目：基于检索增强的电商退换货智能客服系统设计与实现",
            "",
            "学院：人工智能学院",
            "年级：2025届",
            "专业及班级：软件231",
            "组员姓名学号：刘剑宇（学号提交前补充）",
        ]:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if line:
                r = p.add_run(line)
                self.font(r, size=17)
        for _ in range(3):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("提交时间：2026年5月")
        self.font(r, size=17)
        self.doc.add_section(WD_SECTION.NEW_PAGE)

    def add_toc(self) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("目录")
        self.font(r, size=16)
        for item in ["一、课题介绍", "二、团队分工", "三、系统设计", "四、实现情况", "五、未来与展望", "评分表"]:
            p = self.doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(item)
            self.font(r, size=12)
        self.doc.add_section(WD_SECTION.NEW_PAGE)

    def h1(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        self.font(r, size=14, east="楷体")

    def h2(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        self.font(r, size=13, east="楷体")

    def h3(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        self.font(r, size=12, bold=True)

    def para(self, text: str) -> None:
        if not text.strip():
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text.strip())
        self.font(r, size=12)

    def add_markdown(self, text: str) -> None:
        for block in re.split(r"\n{2,}", text):
            line = block.strip()
            if not line:
                continue
            if line.startswith("# "):
                self.h1(line[2:].strip())
            elif line.startswith("## "):
                self.h2(line[3:].strip())
            elif line.startswith("### "):
                self.h3(line[4:].strip())
            elif line.startswith("**") and line.endswith("**"):
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(line.strip("*"))
                self.font(r, size=10.5, bold=True)
            elif line.startswith("|"):
                self.markdown_table(line)
            elif line.startswith("- "):
                self.para(line[2:].strip())
            else:
                self.para(line.replace("\n", ""))

    def markdown_table(self, block: str) -> None:
        rows: list[list[str]] = []
        for raw in block.splitlines():
            stripped = raw.strip()
            if not stripped.startswith("|"):
                continue
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if cells and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
                continue
            rows.append(cells)
        if not rows:
            return
        header, body = rows[0], rows[1:]
        self.table(header, body)

    def table(self, headers: list[str], rows: list[list[str]], title: str | None = None) -> None:
        if title:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(title)
            self.font(r, size=12, bold=True)
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            self.set_cell(table.rows[0].cells[i], header, True)
        for row in rows:
            cells = table.add_row().cells
            for i, item in enumerate(row):
                self.set_cell(cells[i], item)
        self.doc.add_paragraph()

    def set_cell(self, cell, text: str, bold: bool = False) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        self.font(r, size=10.5, bold=bold)

    def add_figures(self) -> None:
        self.h2("4.13 真实系统截图")
        self.para("以下截图均来自本项目真实运行页面或项目已有图示。本次修订未使用 AI 生成图片，也未使用其他图像生成模型；若后续确需生成或编辑图像，按 AGENTS.md 规定只能使用 gpt-image-2。")
        for filename, title, note in REAL_FIGURES:
            path = ASSET_DIR / filename
            if not path.exists():
                continue
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(path), width=Cm(15.2))
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(f"图 {self.figure_no} {title}")
            self.font(r, size=9.5, bold=True)
            self.figure_no += 1
            explain = self.doc.add_paragraph()
            explain.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = explain.add_run(note)
            self.font(r, size=9)

    def add_score_table(self) -> None:
        self.table(
            ["模块", "项目", "要求", "分值", "得分"],
            [
                ["开题报告", "内容", "团队分工合理；方案设计可行", "10", ""],
                ["", "质量", "报告逻辑结构清晰；具有一定创新性", "10", ""],
                ["结项答辩", "团队", "满足设计需求；运行顺畅", "30", ""],
                ["", "", "程序设计规范；UI设计友好", "10", ""],
                ["", "个人", "对自身开发部分自述内容准确，能够根据设计要点进行全面阐述", "10", ""],
                ["", "", "回答问题思维敏捷、逻辑清晰、语言流畅、正确合理、有自己的见解", "10", ""],
                ["结项报告", "内容", "内容详实、表达正确、逻辑清晰", "10", ""],
                ["", "拓展", "对本项目提出合理、可行建议", "10", ""],
                ["统计", "难度系数", "选题自拟则难度系数为0.8，否则为1", "", ""],
                ["", "总分", "各项目累积分值", "100", ""],
            ],
            "评分表",
        )


def set_cell_width(cell, width_cm: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        from docx.oxml import OxmlElement

        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def main() -> None:
    for name in REMOVED_GENERATED_FIGURES:
        path = ASSET_DIR / name
        if path.exists():
            path.unlink()
    recovered = recover_markdown_from_head()
    text = clean_markdown(recovered)
    OUT_MD.write_text(text, encoding="utf-8")
    doc = TemplateDoc()
    doc.add_cover()
    doc.add_toc()
    doc.add_markdown(text)
    doc.add_figures()
    doc.add_score_table()
    doc.doc.save(OUT_DOCX)
    chars = sum(1 for c in text if 0x4E00 <= ord(c) <= 0x9FFF)
    print(f"DOCX: {OUT_DOCX}")
    print(f"MD: {OUT_MD}")
    print(f"Chinese chars in markdown: {chars}")
    print(f"Figures: {len([p for p in ASSET_DIR.glob('*.png')])}")
    if chars < 20000:
        raise SystemExit(f"Report is too short: {chars} Chinese chars")


if __name__ == "__main__":
    main()
