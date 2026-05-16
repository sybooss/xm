from __future__ import annotations

import math
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SOURCE_MD = DOCS / "结项报告_刘剑宇个人版.md"
SOURCE_DOCX = DOCS / "结项报告_刘剑宇个人版.docx"
OUT_DOCX = DOCS / "结项报告_刘剑宇个人版_图表数据库增强版.docx"
OUT_MD = DOCS / "结项报告_刘剑宇个人版_图表数据库增强版.md"
REPORT_FIGURE_DIR = ROOT / "assets" / "report_figures"
SOURCE_FIGURE_DIR = DOCS / "final-report-assets" / "liujianyu"
SUMMARY_PATH = REPORT_FIGURE_DIR / "enhancement_summary.md"
QUALITY_PATH = REPORT_FIGURE_DIR / "quality_checklist.md"

FONT_CANDIDATES = [
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
]


@dataclass(frozen=True)
class ReportFigure:
    stem: str
    title: str
    kind: str
    mmd: str
    note: str


@dataclass(frozen=True)
class ScreenshotFigure:
    filename: str
    title: str
    note: str


@dataclass(frozen=True)
class ChapterFigure:
    stem: str
    filename: str
    heading: str
    title: str
    note: str
    entry: tuple[str, ...] = ()
    flow: tuple[str, ...] = ()
    data: tuple[str, ...] = ()
    checks: tuple[str, ...] = ()


ENGINEERING_FIGURES = [
    ReportFigure(
        "fig_a_system_architecture",
        "系统总体技术架构图",
        "layers",
        """flowchart LR
    User[顾客端 / 管理员端]
    Vue[Vue 3 + Element Plus]
    API[Axios API 封装 / JWT]
    Spring[Spring Boot Controller]
    Service[Service 业务编排]
    Mapper[MyBatis Mapper]
    MySQL[(MySQL 8)]
    AI[LangChain4j / OpenAI-Compatible]
    KB[knowledge_doc 知识库]
    Logs[retrieval_log / ai_call_log / process_trace]
    User --> Vue --> API --> Spring --> Service --> Mapper --> MySQL
    Service --> AI
    Service --> KB
    Service --> Logs
    KB --> MySQL
    Logs --> MySQL
""",
        "该图说明系统不是单一聊天页面，而是由 Vue 前端、Spring Boot 服务层、MyBatis 数据访问、MySQL 数据模型和 LangChain4j AI 增强层共同组成。AI 位于辅助层，不绕过 Service 层业务校验。",
    ),
    ReportFigure(
        "fig_b_database_er",
        "数据库核心实体关系图",
        "er",
        """erDiagram
    USER_ACCOUNT ||--o{ DEMO_ORDER : owns
    USER_ACCOUNT ||--o{ AFTER_SALE_APPLICATION : submits
    USER_ACCOUNT ||--o{ SERVICE_REVIEW : writes
    DEMO_ORDER ||--o{ AFTER_SALE_APPLICATION : has
    AFTER_SALE_APPLICATION ||--o{ AFTER_SALE_PROCESS_LOG : records
    AFTER_SALE_APPLICATION ||--o{ AFTER_SALE_EVIDENCE : contains
    AFTER_SALE_APPLICATION ||--o{ REPLY_DRAFT : generates
    AFTER_SALE_APPLICATION ||--o{ AFTER_SALE_RISK_ASSESSMENT : evaluated_by
    AFTER_SALE_EVIDENCE ||--o{ EVIDENCE_AUDIT : audited_by
    CHAT_SESSION ||--o{ SERVICE_TICKET : creates
    SERVICE_TICKET ||--o{ REPLY_DRAFT : assists
    KNOWLEDGE_DOC ||--o{ RETRIEVAL_LOG : hit_by
    CHAT_SESSION ||--o{ AI_CALL_LOG : traces
    CHAT_SESSION ||--o{ PROCESS_TRACE : records
""",
        "该图以订单和售后申请为中心，展示顾客、订单、售后、凭证、草稿、风险评估、评价、知识检索和日志之间的关系。service_ticket 在当前 SQL 中主要通过会话、订单、用户关联，并由 after_sale_application.ticket_id 建立业务回指。",
    ),
    ReportFigure(
        "fig_c_after_sale_state_machine",
        "售后状态机图",
        "state",
        """stateDiagram-v2
    [*] --> SUBMITTED: 顾客提交售后
    SUBMITTED --> UNDER_REVIEW: 管理员开始审核
    UNDER_REVIEW --> NEED_MORE_EVIDENCE: 要求补充材料
    NEED_MORE_EVIDENCE --> UNDER_REVIEW: 顾客补充凭证
    UNDER_REVIEW --> APPROVED: 审核通过
    UNDER_REVIEW --> REJECTED: 审核驳回
    APPROVED --> COMPLETED: 管理员确认完成
    UNDER_REVIEW --> TICKET_CREATED: 创建人工工单
    NEED_MORE_EVIDENCE --> TICKET_CREATED: 复杂问题转人工
    TICKET_CREATED --> UNDER_REVIEW: 工单处理后继续审核
    COMPLETED --> REVIEWED: 顾客评价
""",
        "该图强调售后单不是普通 CRUD 数据，而是需要受状态机约束。每一次状态变化都应伴随 after_sale_process_log 记录，便于答辩时追溯处理依据。",
    ),
    ReportFigure(
        "fig_d_submit_after_sale_sequence",
        "顾客提交售后到管理员审核时序图",
        "sequence",
        """sequenceDiagram
    actor Customer as 顾客
    participant Vue as Vue 顾客端页面
    participant API as Axios/API 封装
    participant Controller as CustomerAfterSaleController
    participant Service as AfterSaleApplicationService
    participant Mapper as MyBatis Mapper
    participant DB as MySQL
    Customer->>Vue: 填写售后类型、原因、金额和说明
    Vue->>API: POST /customer/after-sales
    API->>Controller: 携带 JWT 与请求体
    Controller->>Service: submit(request, currentUser)
    Service->>DB: 校验订单归属与状态
    Service->>Mapper: 写入 after_sale_application
    Service->>Mapper: 写入 after_sale_process_log
    Service->>DB: 更新订单售后状态
    Service-->>Controller: 返回售后详情
    Controller-->>Vue: Result.success(data)
    Vue-->>Customer: 刷新列表和时间线
""",
        "该图把一次顾客提交售后的完整后端链路展开，说明页面动作、接口、Service 校验、Mapper 落库和日志记录之间形成闭环。",
    ),
    ReportFigure(
        "fig_e_reply_draft_audit_flow",
        "AI 草稿生成与审计链路图",
        "flow",
        """flowchart LR
    Admin[管理员审核台]
    Context[订单 / 售后 / 凭证 / 日志]
    Knowledge[知识检索]
    Template[本地模板草稿]
    AI[LangChain4j 表达增强]
    Draft[reply_draft 保存草稿]
    Log[retrieval_log / ai_call_log / process_log]
    Action[采纳或废弃]
    Admin --> Context --> Knowledge --> Template --> AI --> Draft --> Action
    Knowledge --> Log
    AI --> Log
    Action --> Log
""",
        "该图说明 AI 只参与回复草稿生成和表达增强，管理员采纳或废弃后留下审计记录，AI 不直接决定售后状态。",
    ),
    ReportFigure(
        "fig_f_image_risk_c2pa_flow",
        "图片风险与 C2PA 预审数据流图",
        "flow",
        """flowchart LR
    Upload[图片上传]
    Message[聊天消息 / 凭证记录]
    Risk[图片风险扫描]
    C2PA[C2PA 内容凭证检测]
    Audit[evidence_audit 风险结果]
    Panel[聊天面板风险展示]
    Trace[process_trace 处理轨迹]
    Manual[人工复核 / 补证建议]
    Upload --> Message --> Risk --> C2PA --> Audit --> Panel --> Manual
    Risk --> Trace
    Audit --> Trace
""",
        "该图说明图片风险与 C2PA 检测只提供可信度预审信号。系统不会把未发现 C2PA 凭证写成图片造假结论，而是结合人工复核和补证建议处理。",
    ),
    ReportFigure(
        "fig_g_sla_ticket_flow",
        "SLA 与人工工单协同流程图",
        "flow",
        """flowchart LR
    Application[售后申请]
    Risk[风险评估 / 优先级]
    SLA[SLA 队列]
    Review[管理员审核台]
    Ticket[人工工单]
    Follow[客服跟进]
    Log[after_sale_process_log 回写]
    Application --> Risk --> SLA --> Review --> Ticket --> Follow --> Log
    Review --> Log
    Ticket --> Review
""",
        "该图说明 SLA 队列、风险识别、人工工单和审核台之间的协作关系，突出复杂售后可以被人工接管并回写处理日志。",
    ),
]

SCREENSHOT_FIGURES = [
    ScreenshotFigure("login.png", "登录与角色入口", "登录页区分管理员和顾客入口，是双端权限边界的第一层。"),
    ScreenshotFigure("customer.png", "顾客端我的售后", "顾客可以查看订单、提交售后、补充凭证、跟踪进度并评价服务。"),
    ScreenshotFigure("admin_review.png", "管理员售后审核工作台", "管理员在同一页面处理申请、证据、决策、草稿和审计。"),
    ScreenshotFigure("sla.png", "SLA 跟进中心", "SLA 中心将超时、优先级和待补材料售后聚合成风险任务队列。"),
    ScreenshotFigure("ticket.png", "人工工单页面", "投诉、物流异常或复杂售后可升级为人工客服工单并持续跟进。"),
    ScreenshotFigure("chat_workbench.png", "咨询接待工作台", "咨询工作台展示会话、订单上下文、AI 建议、知识依据和处理轨迹。"),
    ScreenshotFigure("chat_risk.png", "图片风险与 C2PA 预审", "聊天图片会触发真实性预审，展示 AI 生成、水印、篡改和 C2PA 信号。"),
    ScreenshotFigure("knowledge.png", "知识库管理与检索", "知识库支撑 RAG 依据展示和回复草稿中的规则引用。"),
    ScreenshotFigure("orders.png", "订单与售后上下文", "订单管理页面保存商品、物流、支付和售后状态，是业务判断依据。"),
    ScreenshotFigure("logs.png", "服务日志诊断", "日志中心集中呈现 AI 调用、知识检索和处理轨迹，支撑可解释性。"),
    ScreenshotFigure("ai_test.png", "AI 质检页面", "AI 质检页面验证当前 LangChain4j 模型链路和本地兜底状态。"),
]

CHAPTER4_SECTION_FIGURES = [
    ChapterFigure("fig_4_01_login_permission", "fig_4_01_login_permission.png", "4.1 登录注册与双端权限控制", "登录注册与双端权限控制", "该图说明顾客端、管理员端、JWT 令牌、路由守卫、后端拦截器和 user_account 表之间的关系。", ("顾客登录入口", "管理员登录入口", "前端路由守卫"), ("提交账号密码", "后端校验用户", "签发 JWT 令牌", "按角色进入页面"), ("user_account", "JWT 拦截器", "角色权限字段"), ("无令牌禁止访问", "角色不匹配禁止越权", "登录状态可追溯")),
    ChapterFigure("fig_4_02_customer_after_sale", "fig_4_02_customer_after_sale.png", "4.2 顾客端我的售后", "顾客端我的售后闭环", "该图展示顾客端从订单入口到售后申请、凭证补充、进度查看和服务评价的闭环。", ("我的订单", "申请售后", "补充凭证", "服务评价"), ("选择订单", "填写售后原因", "上传凭证材料", "查看处理进度", "提交服务评价"), ("demo_order", "after_sale_application", "after_sale_evidence", "service_review"), ("订单归属校验", "补证入口按状态展示", "评价只在完成后开放")),
    ChapterFigure("fig_4_03_admin_review", "fig_4_03_admin_review.png", "4.3 管理员售后审核工作台", "管理员售后审核工作台", "该图体现审核列表、详情抽屉、状态变更和 after_sale_process_log 审计记录。", ("审核列表", "售后详情", "凭证查看", "处理按钮"), ("筛选待审申请", "查看订单与凭证", "填写处理意见", "变更售后状态", "写入过程日志"), ("after_sale_application", "after_sale_process_log", "after_sale_evidence"), ("状态流转受控", "处理意见留痕", "顾客端同步结果")),
    ChapterFigure("fig_4_04_chat_workbench", "fig_4_04_chat_workbench.png", "4.4 咨询工作台与多轮客服链路", "咨询工作台与多轮客服链路", "该图展示 chat_session、chat_message、上下文检索、AI 建议和人工接待协同。", ("会话列表", "消息输入区", "订单上下文", "处理轨迹"), ("顾客发起咨询", "读取会话历史", "匹配订单和知识", "生成回复建议", "必要时转人工"), ("chat_session", "chat_message", "knowledge_doc", "process_trace"), ("多轮上下文连续", "回复依据可查看", "人工接管不丢失记录")),
    ChapterFigure("fig_4_05_knowledge_rag", "fig_4_05_knowledge_rag.png", "4.5 知识库管理与 RAG 依据展示", "知识库管理与 RAG 依据展示", "该图说明 knowledge_doc、retrieval_log、规则依据和回复草稿引用之间的关系。", ("知识文档维护", "检索入口", "依据展示", "草稿引用"), ("维护售后规则", "按问题检索知识", "记录命中文档", "在回复中引用依据"), ("knowledge_doc", "retrieval_log", "reply_draft.knowledge_refs"), ("规则来源清楚", "命中记录可追踪", "回复不脱离业务规则")),
    ChapterFigure("fig_4_06_order_context", "fig_4_06_order_context.png", "4.6 订单管理与售后上下文", "订单管理与售后上下文", "该图突出 demo_order、订单状态、物流支付信息和售后判断依据。", ("订单列表", "订单详情", "售后入口", "物流信息"), ("读取顾客订单", "判断订单状态", "关联售后申请", "形成处理依据"), ("demo_order", "after_sale_application", "user_account"), ("只能处理本人订单", "退款换货依赖订单状态", "订单与售后保持一致")),
    ChapterFigure("fig_4_07_sla_queue", "fig_4_07_sla_queue.png", "4.7 SLA 跟进与风险任务队列", "SLA 跟进与风险任务队列", "该图展示超时、优先级、风险标记和待处理队列的组织方式。", ("SLA 中心", "风险筛选", "优先级标记", "待处理队列"), ("汇总售后任务", "识别超时风险", "按优先级排序", "推送管理员处理"), ("after_sale_application", "after_sale_risk_assessment", "service_ticket"), ("超时任务可定位", "高风险任务优先", "处理后状态回写")),
    ChapterFigure("fig_4_08_ticket_collaboration", "fig_4_08_ticket_collaboration.png", "4.8 人工工单协同", "人工工单协同", "该图体现 service_ticket、客服分派、处理回写和售后联动。", ("工单列表", "分派客服", "处理记录", "售后联动"), ("创建人工工单", "指定处理人员", "记录跟进内容", "回写售后单状态"), ("service_ticket", "after_sale_application.ticket_id", "after_sale_process_log"), ("复杂问题可接管", "工单与售后互相回指", "处理过程可复盘")),
    ChapterFigure("fig_4_09_audit_reply_draft", "fig_4_09_audit_reply_draft.png", "4.9 AI 副驾驶可审计回复草稿", "AI 副驾驶可审计回复草稿", "该图说明 reply_draft、知识依据、管理员采纳或废弃和审计留痕。", ("草稿按钮", "知识依据", "采纳/废弃", "审核日志"), ("读取售后上下文", "形成回复草稿", "管理员审核确认", "记录采纳结果"), ("reply_draft", "knowledge_doc", "ai_call_log", "after_sale_process_log"), ("草稿不自动发出", "关键状态仍由管理员确认", "采纳废弃均留痕")),
    ChapterFigure("fig_4_10_image_risk_c2pa", "fig_4_10_image_risk_c2pa.png", "4.10 聊天图片风险扫描与 C2PA/图片可信度检测", "聊天图片风险扫描与 C2PA 可信度检测", "该图强调 image_risk、C2PA 信号、evidence_audit 和人工复核。", ("图片消息", "风险面板", "C2PA 状态", "复核入口"), ("上传聊天图片", "提取可信度信号", "形成风险提示", "必要时转人工复核"), ("chat_message", "image_risk", "evidence_audit", "process_trace"), ("风险信号不等于结论", "缺少凭证不直接判假", "高风险提示补充材料")),
    ChapterFigure("fig_4_11_customer_profile_review", "fig_4_11_customer_profile_review.png", "4.11 客户画像与服务评价", "客户画像与服务评价", "该图展示 customer_profile、service_review、满意度和历史售后信息。", ("客户资料", "历史售后", "满意度评价", "服务标签"), ("汇总顾客记录", "统计售后历史", "保存评价内容", "辅助服务分层"), ("customer_profile", "service_review", "after_sale_application"), ("评价来源可查", "画像不替代审核", "历史记录辅助判断")),
    ChapterFigure("fig_4_12_product_quality_alert", "fig_4_12_product_quality_alert.png", "4.12 商品质量问题聚合预警", "商品质量问题聚合预警", "该图体现商品维度聚合、售后原因、质量问题和运营预警。", ("商品维度", "售后原因", "质量标签", "运营预警"), ("汇总售后原因", "按商品聚合统计", "识别异常集中问题", "提示运营复核"), ("demo_order.product_name", "after_sale_application.reason_text", "after_sale_risk_assessment"), ("只做趋势提示", "数据来自售后记录", "异常商品需要人工确认")),
    ChapterFigure("fig_4_13_log_ai_quality", "fig_4_13_log_ai_quality.png", "4.13 日志诊断与 AI 质检", "日志诊断与 AI 质检", "该图覆盖 ai_call_log、retrieval_log、process_trace 和 AI 质检页面。", ("日志中心", "质检页面", "检索记录", "调用状态"), ("发起测试请求", "记录模型状态", "保存检索命中", "展示失败兜底"), ("ai_call_log", "retrieval_log", "process_trace"), ("接口失败不影响主流程", "日志记录便于排查", "兜底状态可展示")),
    ChapterFigure("fig_4_14_diagnosis_risk_evidence", "fig_4_14_diagnosis_risk_evidence.png", "4.14 售后前置诊断、风险评估与凭证审核", "售后前置诊断、风险评估与凭证审核", "该图展示前置诊断、risk_assessment、after_sale_evidence 和审核建议。", ("售后原因", "凭证材料", "风险评分", "审核建议"), ("读取申请内容", "检查凭证充分性", "计算风险等级", "给出补证建议"), ("after_sale_application", "after_sale_evidence", "after_sale_risk_assessment", "evidence_audit"), ("风险只作辅助", "凭证缺失可补充", "审核结论人工确认")),
    ChapterFigure("fig_4_15_operation_status", "fig_4_15_operation_status.png", "4.15 运营首页与系统状态展示", "运营首页与系统状态展示", "该图汇总工单数量、SLA 风险、系统状态和关键指标。", ("运营首页", "系统状态", "关键指标", "快捷入口"), ("加载统计数据", "展示数据库状态", "展示模型状态", "跳转核心模块"), ("system_status", "after_sale_application", "service_ticket", "ai_call_log"), ("演示前可检查环境", "业务状态和模型状态分开", "入口指向真实模块")),
    ChapterFigure("fig_4_16_dual_after_sale_loop", "fig_4_16_dual_after_sale_loop.png", "4.16 真实双端售后业务闭环", "真实双端售后业务闭环", "该图说明顾客提交、管理员审核、凭证补充和评价回收。", ("顾客端售后", "管理员审核", "补证入口", "服务评价"), ("顾客提交申请", "管理员要求补证", "顾客补充材料", "管理员确认完成", "顾客评价服务"), ("after_sale_application", "after_sale_evidence", "after_sale_process_log", "service_review"), ("双端状态同步", "补证动作可追踪", "完成后形成评价闭环")),
    ChapterFigure("fig_4_17_sla_ticket_collaboration", "fig_4_17_sla_ticket_collaboration.png", "4.17 SLA 跟进与人工工单协同", "SLA 跟进与人工工单协同", "该图展示 service_ticket、客服跟进、SLA 队列和日志回写。", ("SLA 队列", "人工工单", "客服跟进", "日志回写"), ("识别高优先级售后", "创建或关联工单", "客服持续处理", "回写处理结果"), ("service_ticket", "after_sale_application.ticket_id", "after_sale_process_log"), ("队列与工单一致", "处理人员可明确", "售后详情可看到工单")),
    ChapterFigure("fig_4_18_ai_reply_draft", "fig_4_18_ai_reply_draft.png", "4.18 AI 副驾驶回复草稿", "AI 副驾驶回复草稿", "该图强调 LangChain4j、reply_draft、采纳状态和审计字段。", ("草稿生成", "知识依据", "人工确认", "审计字段"), ("读取售后上下文", "结合知识依据", "保存回复草稿", "管理员采纳或废弃"), ("reply_draft", "ai_call_log", "retrieval_log", "process_trace"), ("失败时本地兜底", "草稿不能直接改状态", "采纳结果写入日志")),
    ChapterFigure("fig_4_19_image_risk_detail", "fig_4_19_image_risk_detail.png", "4.19 图片风险扫描与 C2PA 可信度检测", "图片风险扫描与 C2PA 可信度检测", "该图展示图片消息、可信度检测、风险标签和复核入口。", ("图片上传", "风险标签", "元数据线索", "人工复核"), ("保存图片消息", "分析可信度信号", "生成风险标签", "提示补充原始材料"), ("chat_message", "image_risk", "evidence_audit"), ("风险文案保持谨慎", "不把检测结果写成绝对结论", "复核入口保留人工判断")),
    ChapterFigure("fig_4_20_database_consistency", "fig_4_20_database_consistency.png", "4.20 数据库实现与数据一致性验证", "数据库实现与数据一致性验证", "该图覆盖 schema.sql、主外键、索引、事务一致性和初始化数据。", ("schema.sql", "主外键关系", "索引设计", "初始化数据"), ("创建基础表", "建立业务关联", "按场景写入数据", "跨页面验证一致性"), ("user_account", "demo_order", "after_sale_application", "service_ticket", "reply_draft"), ("字段名称按 SQL 校对", "工单通过 ticket_id 回指", "状态变化必须写日志")),
    ChapterFigure("fig_4_21_validation_matrix", "fig_4_21_validation_matrix.png", "4.21 测试与验证", "测试与验证", "该图串联文档生成、接口验证、页面截图和渲染抽查。", ("接口检查", "页面检查", "文档生成", "渲染抽查"), ("执行生成脚本", "统计图表数量", "渲染 DOCX/PDF", "抽查页面版式"), ("增强版报告", "质量检查清单", "渲染页面截图"), ("图数与章节对应", "关键词无制作痕迹", "表格和图片不重叠")),
    ChapterFigure("fig_4_22_project_innovation", "fig_4_22_project_innovation.png", "4.22 项目特色与创新点", "项目特色与创新点", "该图概括业务闭环、AI 辅助、RAG 依据、风险预审和可审计。", ("双端售后闭环", "RAG 依据展示", "风险预审", "日志审计"), ("业务规则先校验", "知识库提供依据", "草稿辅助表达", "人工确认关键动作"), ("knowledge_doc", "reply_draft", "image_risk", "process_trace"), ("业务系统为主", "智能能力为辅", "关键动作有证据链")),
    ChapterFigure("fig_4_23_personal_database_mapping", "fig_4_23_personal_database_mapping.png", "4.23 个人负责模块与数据库支撑关系", "个人负责模块与数据库支撑关系", "该图说明售后、工单、知识库、草稿、日志与数据表的支撑关系。", ("售后闭环", "工单协同", "回复草稿", "图片风险"), ("模块功能落到接口", "接口动作写入数据表", "数据表支撑页面回显", "日志记录处理过程"), ("after_sale_application", "service_ticket", "reply_draft", "image_risk", "process_trace"), ("个人分工边界清楚", "模块之间有数据关联", "答辩可按链路说明")),
    ChapterFigure("fig_4_24_personal_module_summary", "fig_4_24_personal_module_summary.png", "4.24 个人负责模块补充说明", "个人负责模块补充说明", "该图整理模块边界、接口链路、数据落表和答辩要点。", ("模块边界", "接口链路", "数据落表", "答辩要点"), ("说明负责范围", "对应页面操作", "对应后端接口", "对应数据库表", "总结验证证据"), ("售后表", "工单表", "草稿表", "日志表", "风险表"), ("不夸大功能", "强调可运行闭环", "用截图和图示支撑说明")),
]


def font_path() -> Path | None:
    for path in FONT_CANDIDATES:
        if path.exists():
            return path
    return None


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    path = font_path()
    if path:
        return ImageFont.truetype(str(path), size=size, index=0)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        trial = current + char
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = "#1f2937",
) -> None:
    x1, y1, x2, y2 = box
    manual_lines = text.split("\n")
    lines: list[str] = []
    for item in manual_lines:
        lines.extend(wrap_text(draw, item, font, x2 - x1 - 28))
    line_h = max(18, int(font.size * 1.25 if hasattr(font, "size") else 22))
    total_h = line_h * len(lines)
    y = y1 + (y2 - y1 - total_h) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        x = x1 + (x2 - x1 - (bbox[2] - bbox[0])) // 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str = "#ffffff",
    outline: str = "#4b5563",
    width: int = 2,
    font_size: int = 25,
) -> None:
    draw.rectangle(box, fill=fill, outline=outline, width=width)
    draw_centered_text(draw, box, text, load_font(font_size), fill="#111827")


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#374151") -> None:
    draw.line([start, end], fill=fill, width=3)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head = 13
    points = [
        end,
        (int(end[0] - head * math.cos(angle - math.pi / 6)), int(end[1] - head * math.sin(angle - math.pi / 6))),
        (int(end[0] - head * math.cos(angle + math.pi / 6)), int(end[1] - head * math.sin(angle + math.pi / 6))),
    ]
    draw.polygon(points, fill=fill)


def draw_orthogonal_arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    fill: str = "#374151",
) -> None:
    for i in range(max(0, len(points) - 2)):
        draw.line([points[i], points[i + 1]], fill=fill, width=3)
    if len(points) >= 2:
        draw_arrow(draw, points[-2], points[-1], fill)


def draw_header(draw: ImageDraw.ImageDraw, title: str, subtitle: str | None = None) -> None:
    draw.text((70, 48), title, font=load_font(42), fill="#111827")
    draw.line((70, 112, 610, 112), fill="#111827", width=4)
    if subtitle:
        draw.text((70, 132), subtitle, font=load_font(23), fill="#4b5563")


def draw_caption(draw: ImageDraw.ImageDraw, text: str, y: int, width: int = 1600) -> None:
    font = load_font(22)
    lines = wrap_text(draw, text, font, width - 160)
    for idx, line in enumerate(lines):
        draw.text((80, y + idx * 30), line, font=font, fill="#374151")


def draw_small_label(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str) -> None:
    draw.text(xy, text, font=load_font(20), fill="#374151")


def draw_panel_title(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str) -> None:
    x1, y1, x2, _ = box
    draw.rectangle((x1, y1, x2, y1 + 42), fill="#eeeeee", outline="#444444", width=2)
    draw.text((x1 + 18, y1 + 8), title, font=load_font(23), fill="#111111")


def draw_entity(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    fields: list[str],
    fill: str = "#ffffff",
) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill=fill, outline="#333333", width=2)
    draw.rectangle((x1, y1, x2, y1 + 42), fill="#eeeeee", outline="#333333", width=2)
    title_bbox = draw.textbbox((0, 0), title, font=load_font(22))
    draw.text((x1 + (x2 - x1 - (title_bbox[2] - title_bbox[0])) // 2, y1 + 8), title, font=load_font(22), fill="#111111")
    row_h = max(28, (y2 - y1 - 42) // max(1, len(fields)))
    y = y1 + 42
    for field in fields:
        draw.line((x1, y, x2, y), fill="#bdbdbd", width=1)
        draw.text((x1 + 12, y + 5), field, font=load_font(17), fill="#222222")
        y += row_h


def draw_swimlane(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill="#fafafa", outline="#555555", width=2)
    draw.rectangle((x1, y1, x1 + 92, y2), fill="#eeeeee", outline="#555555", width=2)
    chars = list(title)
    top = y1 + max(18, (y2 - y1 - len(chars) * 26) // 2)
    for idx, char in enumerate(chars):
        draw.text((x1 + 34, top + idx * 26), char, font=load_font(22), fill="#111111")


def save_image(image: Image.Image, stem: str) -> Path:
    png = REPORT_FIGURE_DIR / f"{stem}.png"
    image.save(png)
    return png


def render_layer_figure(fig: ReportFigure) -> Path:
    image = Image.new("RGB", (1600, 900), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw_header(draw, fig.title, "Spring Boot + Vue 3 + MySQL + LangChain4j 的分层闭环")
    layers = [
        ("表示层", "顾客端 / 管理员端\nVue 3 + Element Plus", 190),
        ("接口层", "Axios 请求封装 / JWT 身份令牌\nREST Controller 入参校验", 300),
        ("业务层", "AfterSaleService / TicketService / AiService\n状态流转、风控、补证、工单与回复草稿规则", 410),
        ("数据访问层", "MyBatis Mapper\n订单、售后、凭证、日志、知识库映射", 520),
        ("数据层", "MySQL 8\nschema.sql 约束、索引与业务表", 630),
    ]
    for index, (name, label, y) in enumerate(layers):
        draw.rectangle((90, y, 1320, y + 78), fill="#f7f7f7" if index % 2 else "#ffffff", outline="#333333", width=2)
        draw.rectangle((90, y, 260, y + 78), fill="#eeeeee", outline="#333333", width=2)
        draw_centered_text(draw, (90, y, 260, y + 78), name, load_font(26), "#111111")
        draw_centered_text(draw, (285, y, 1300, y + 78), label, load_font(24), "#111111")
        if index:
            draw_arrow(draw, (705, layers[index - 1][2] + 78), (705, y), "#333333")
    side_boxes = [
        ((1370, 230, 1530, 330), "知识库\nknowledge_doc"),
        ((1370, 380, 1530, 480), "AI 增强\nLangChain4j"),
        ((1370, 530, 1530, 630), "审计日志\n三类日志"),
    ]
    for box, label in side_boxes:
        draw_box(draw, box, label, "#ffffff", "#333333", 2, 22)
    draw_orthogonal_arrow(draw, [(1320, 448), (1345, 448), (1345, 280), (1370, 280)], "#333333")
    draw_orthogonal_arrow(draw, [(1320, 448), (1370, 430)], "#333333")
    draw_orthogonal_arrow(draw, [(1320, 448), (1345, 448), (1345, 580), (1370, 580)], "#333333")
    draw_caption(draw, "设计要点：AI 只提供检索、改写和风险提示，不绕过 Service 层；订单、售后、凭证、工单和日志均由业务规则校验后落库。", 760)
    return save_image(image, fig.stem)


def render_er_figure(fig: ReportFigure) -> Path:
    image = Image.new("RGB", (1600, 1050), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw_header(draw, fig.title, "围绕 after_sale_application 的主从表与日志表关系")
    entities: dict[str, tuple[tuple[int, int, int, int], str, list[str]]] = {
        "user": ((80, 165, 350, 330), "user_account", ["PK id", "username", "role", "status"]),
        "order": ((430, 165, 700, 330), "demo_order", ["PK id", "order_no", "user_id", "amount"]),
        "after": ((790, 145, 1150, 360), "after_sale_application", ["PK id", "order_id / user_id", "type / reason / status", "ticket_id", "created_at"]),
        "review": ((1240, 165, 1510, 330), "service_review", ["PK id", "application_id", "rating", "content"]),
        "log": ((95, 450, 395, 610), "after_sale_process_log", ["PK id", "application_id", "operator_id", "action / remark"]),
        "evidence": ((485, 450, 785, 610), "after_sale_evidence", ["PK id", "application_id", "file_url", "audit_status"]),
        "draft": ((875, 450, 1175, 610), "reply_draft", ["PK id", "application_id", "source", "adopted"]),
        "ticket": ((1265, 450, 1535, 610), "service_ticket", ["PK id", "application_id", "priority", "owner_id"]),
        "knowledge": ((175, 745, 475, 905), "knowledge_doc", ["PK id", "title", "category", "content"]),
        "retrieval": ((560, 745, 860, 905), "retrieval_log", ["PK id", "doc_id", "session_id", "score"]),
        "ai": ((945, 745, 1245, 905), "ai_call_log", ["PK id", "application_id", "model", "latency"]),
        "trace": ((1305, 745, 1545, 905), "process_trace", ["PK id", "biz_type", "biz_id", "detail"]),
    }
    for key, (box, title, fields) in entities.items():
        fill = "#f7f7f7" if key == "after" else "#ffffff"
        draw_entity(draw, box, title, fields, fill)

    def anchor(name: str, side: str, offset: int = 0) -> tuple[int, int]:
        x1, y1, x2, y2 = entities[name][0]
        if side == "left":
            return (x1, (y1 + y2) // 2 + offset)
        if side == "right":
            return (x2, (y1 + y2) // 2 + offset)
        if side == "top":
            return ((x1 + x2) // 2 + offset, y1)
        return ((x1 + x2) // 2 + offset, y2)

    def relation(points: list[tuple[int, int]], label: str, label_xy: tuple[int, int]) -> None:
        draw.line(points, fill="#333333", width=2, joint="curve")
        draw.rectangle((label_xy[0] - 6, label_xy[1] - 4, label_xy[0] + 52, label_xy[1] + 24), fill="#ffffff")
        draw.text(label_xy, label, font=load_font(18), fill="#111111")

    relation([anchor("user", "right"), anchor("order", "left")], "1:N", (370, 235))
    relation([anchor("order", "right"), anchor("after", "left")], "1:N", (725, 235))
    relation([anchor("after", "right"), anchor("review", "left")], "1:1", (1160, 235))
    relation([anchor("after", "bottom", -110), anchor("log", "top")], "1:N", (570, 395))
    relation([anchor("after", "bottom", -35), anchor("evidence", "top")], "1:N", (715, 395))
    relation([anchor("after", "bottom", 55), anchor("draft", "top")], "1:N", (1010, 395))
    relation([anchor("after", "bottom", 130), anchor("ticket", "top")], "1:0..1", (1210, 395))
    relation([anchor("knowledge", "right"), anchor("retrieval", "left")], "1:N", (495, 815))
    relation([anchor("retrieval", "right"), anchor("ai", "left")], "N:1", (885, 815))
    relation([anchor("ai", "right"), anchor("trace", "left")], "N:1", (1265, 815))
    draw_caption(draw, "注：该图保留数据库核心链路，重点体现售后主表与凭证、回复草稿、人工工单、知识检索和审计日志之间的约束关系。", 970)
    return save_image(image, fig.stem)


def render_state_figure(fig: ReportFigure) -> Path:
    image = Image.new("RGB", (1500, 900), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw_header(draw, fig.title, "售后主状态、补证分支、驳回分支与转人工分支")
    coords = {
        "SUBMITTED\n已提交": (90, 235, 300, 335),
        "UNDER_REVIEW\n审核中": (405, 235, 645, 335),
        "APPROVED\n已通过": (780, 235, 1000, 335),
        "COMPLETED\n已完成": (1135, 235, 1345, 335),
        "REVIEWED\n已评价": (1135, 515, 1345, 615),
        "NEED_MORE_EVIDENCE\n待补证": (405, 515, 700, 615),
        "REJECTED\n已驳回": (780, 515, 1000, 615),
        "TICKET_CREATED\n已转人工": (695, 690, 1015, 790),
    }
    for label, box in coords.items():
        draw_box(draw, box, label, "#ffffff", "#333333", 2, 23)

    def right(name: str) -> tuple[int, int]:
        x1, y1, x2, y2 = coords[name]
        return (x2, (y1 + y2) // 2)

    def left(name: str) -> tuple[int, int]:
        x1, y1, x2, y2 = coords[name]
        return (x1, (y1 + y2) // 2)

    def top(name: str) -> tuple[int, int]:
        x1, y1, x2, _ = coords[name]
        return ((x1 + x2) // 2, y1)

    def bottom(name: str) -> tuple[int, int]:
        x1, _, x2, y2 = coords[name]
        return ((x1 + x2) // 2, y2)

    def label(text: str, point: tuple[int, int]) -> None:
        draw.text(point, text, font=load_font(18), fill="#111111")

    draw_arrow(draw, right("SUBMITTED\n已提交"), left("UNDER_REVIEW\n审核中"), "#333333")
    label("提交申请", (320, 260))
    draw_arrow(draw, right("UNDER_REVIEW\n审核中"), left("APPROVED\n已通过"), "#333333")
    label("审核通过", (670, 260))
    draw_arrow(draw, right("APPROVED\n已通过"), left("COMPLETED\n已完成"), "#333333")
    label("退款/换货完成", (1030, 260))
    draw_arrow(draw, bottom("COMPLETED\n已完成"), top("REVIEWED\n已评价"), "#333333")
    label("用户评价", (1190, 400))
    draw_orthogonal_arrow(draw, [(525, 335), (525, 515)], "#333333")
    label("要求补证", (545, 410))
    draw_orthogonal_arrow(draw, [(405, 565), (330, 565), (330, 285), (405, 285)], "#333333")
    label("补证后复审", (200, 420))
    draw_orthogonal_arrow(draw, [(645, 305), (890, 455), (890, 515)], "#333333")
    label("材料不通过", (640, 420))
    draw_orthogonal_arrow(draw, [(552, 615), (552, 740), (695, 740)], "#333333")
    label("争议升级", (565, 650))
    draw_orthogonal_arrow(draw, [(765, 690), (765, 650), (645, 650), (645, 335)], "#333333")
    label("人工处理后回到审核", (795, 640))
    draw.line((70, 845, 1430, 845), fill="#cccccc", width=1)
    draw_caption(draw, "状态变化要求：修改 after_sale_application.status 时，同步写入 after_sale_process_log；争议、超时或高风险申请可转入 service_ticket 人工处理。", 805, 1500)
    return save_image(image, fig.stem)


def render_linear_flow(fig: ReportFigure) -> Path:
    image = Image.new("RGB", (1600, 850), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw_header(draw, fig.title, "接口调用、业务校验与数据落库路径")
    if "submit" in fig.stem:
        steps = [
            ("01", "顾客填写申请", "customer 页面提交类型、原因、凭证"),
            ("02", "接口鉴权", "Axios 携带 JWT 调用售后接口"),
            ("03", "业务校验", "Service 校验订单归属、状态和凭证"),
            ("04", "数据落库", "写入 after_sale_application 与日志"),
            ("05", "页面刷新", "返回最新时间线和审核状态"),
        ]
        side_notes = [("after_sale_application", "售后主表"), ("after_sale_process_log", "处理日志"), ("demo_order", "订单状态联动")]
    elif "reply" in fig.stem:
        steps = [
            ("01", "选择售后单", "审核台读取订单与售后上下文"),
            ("02", "知识检索", "查询 knowledge_doc 形成依据"),
            ("03", "本地模板", "先生成可兜底的标准回复"),
            ("04", "AI 增强", "LangChain4j 仅做表达优化"),
            ("05", "草稿审计", "保存 reply_draft 并记录采纳状态"),
        ]
        side_notes = [("retrieval_log", "检索记录"), ("ai_call_log", "AI 调用记录"), ("reply_draft", "回复草稿")]
    elif "image" in fig.stem:
        steps = [
            ("01", "图片接收", "聊天图片或补充凭证进入系统"),
            ("02", "基础校验", "校验格式、大小与业务归属"),
            ("03", "C2PA 检测", "读取元数据和内容来源标记"),
            ("04", "风险入库", "写入 evidence_audit 和风险标签"),
            ("05", "人工复核", "高风险凭证进入人工处理"),
        ]
        side_notes = [("evidence_audit", "凭证审核"), ("process_trace", "风险步骤"), ("service_ticket", "人工复核")]
    elif "sla" in fig.stem:
        steps = [
            ("01", "申请进入队列", "售后单进入审核池"),
            ("02", "风险评估", "按金额、时长和图片风险打标"),
            ("03", "SLA 分级", "按处理时限形成优先级"),
            ("04", "生成工单", "超时或高风险转 service_ticket"),
            ("05", "结果回写", "客服处理后同步状态和日志"),
        ]
        side_notes = [("risk_assessment", "风险评估"), ("service_ticket", "人工工单"), ("process_log", "处理日志")]
    else:
        steps = [("01", "开始", "输入"), ("02", "处理", "业务规则"), ("03", "落库", "数据记录"), ("04", "展示", "页面反馈")]
        side_notes = []

    start_x = 82
    step_w = 236
    step_h = 150
    y = 285
    gap = 38
    boxes: list[tuple[int, int, int, int]] = []
    for idx, (no, title, desc) in enumerate(steps):
        x = start_x + idx * (step_w + gap)
        box = (x, y, x + step_w, y + step_h)
        boxes.append(box)
        draw.rectangle(box, fill="#ffffff", outline="#333333", width=2)
        draw.rectangle((x, y, x + 54, y + step_h), fill="#eeeeee", outline="#333333", width=2)
        no_font = load_font(24)
        no_bbox = draw.textbbox((0, 0), no, font=no_font)
        draw.text((x + (54 - (no_bbox[2] - no_bbox[0])) // 2, y + (step_h - (no_bbox[3] - no_bbox[1])) // 2 - 2), no, font=no_font, fill="#111111")
        draw.text((x + 70, y + 26), title, font=load_font(23), fill="#111111")
        for line_idx, line in enumerate(wrap_text(draw, desc, load_font(18), step_w - 82)):
            draw.text((x + 70, y + 70 + line_idx * 25), line, font=load_font(18), fill="#374151")
        if idx:
            draw_arrow(draw, (boxes[idx - 1][2], y + step_h // 2), (box[0], y + step_h // 2), "#333333")

    if side_notes:
        draw.rectangle((190, 535, 1410, 645), fill="#fafafa", outline="#555555", width=2)
        draw.text((220, 558), "关键落库/审计对象", font=load_font(23), fill="#111111")
        note_x = 500
        for table_name, desc in side_notes:
            draw_box(draw, (note_x, 555, note_x + 220, 625), f"{table_name}\n{desc}", "#ffffff", "#333333", 2, 18)
            note_x += 250
    draw_caption(draw, fig.note, 725)
    return save_image(image, fig.stem)


def render_figure(fig: ReportFigure) -> Path:
    (REPORT_FIGURE_DIR / f"{fig.stem}.mmd").write_text(fig.mmd.strip() + "\n", encoding="utf-8")
    if fig.kind == "layers":
        return render_layer_figure(fig)
    if fig.kind == "er":
        return render_er_figure(fig)
    if fig.kind == "state":
        return render_state_figure(fig)
    return render_linear_flow(fig)


def draw_list_panel(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    items: tuple[str, ...],
) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill="#ffffff", outline="#333333", width=2)
    draw.rectangle((x1, y1, x2, y1 + 48), fill="#eeeeee", outline="#333333", width=2)
    draw.text((x1 + 18, y1 + 12), title, font=load_font(23), fill="#111111")
    y = y1 + 70
    for idx, item in enumerate(items):
        if y + 32 > y2 - 16:
            break
        draw.ellipse((x1 + 22, y + 7, x1 + 34, y + 19), fill="#555555")
        for line in wrap_text(draw, item, load_font(19), x2 - x1 - 62)[:2]:
            draw.text((x1 + 46, y), line, font=load_font(19), fill="#222222")
            y += 26
        if idx < len(items) - 1:
            y += 8


def render_chapter_figure(fig: ChapterFigure) -> Path:
    image = Image.new("RGB", (1600, 980), "#ffffff")
    draw = ImageDraw.Draw(image)
    draw_header(draw, fig.title, "模块流程与数据关系示意")

    draw_list_panel(draw, (75, 190, 440, 430), "页面入口", fig.entry)
    draw_list_panel(draw, (1160, 190, 1525, 430), "验证要点", fig.checks)
    draw_list_panel(draw, (75, 610, 440, 850), "数据支撑", fig.data)

    flow_y = 485
    flow_items = fig.flow or ("页面操作", "业务校验", "数据落库", "结果回显")
    box_w = 220 if len(flow_items) <= 5 else 190
    gap = 30 if len(flow_items) <= 5 else 22
    total_w = len(flow_items) * box_w + (len(flow_items) - 1) * gap
    x = 800 - total_w // 2
    flow_boxes: list[tuple[int, int, int, int]] = []
    for idx, item in enumerate(flow_items):
        box = (x + idx * (box_w + gap), flow_y, x + idx * (box_w + gap) + box_w, flow_y + 96)
        flow_boxes.append(box)
        draw.rectangle(box, fill="#fafafa", outline="#333333", width=2)
        draw.rectangle((box[0], box[1], box[0] + 38, box[3]), fill="#eeeeee", outline="#333333", width=2)
        draw_centered_text(draw, (box[0], box[1], box[0] + 38, box[3]), str(idx + 1), load_font(22), "#111111")
        draw_centered_text(draw, (box[0] + 42, box[1] + 8, box[2] - 8, box[3] - 8), item, load_font(21), "#111111")
        if idx:
            draw_arrow(draw, (flow_boxes[idx - 1][2], flow_y + 48), (box[0], flow_y + 48), "#333333")

    center_box = (560, 190, 1040, 370)
    draw.rectangle(center_box, fill="#f7f7f7", outline="#111111", width=3)
    draw_centered_text(draw, center_box, fig.heading + "\n实现链路", load_font(27), "#111111")
    draw_arrow(draw, (440, 310), (560, 280), "#333333")
    draw_arrow(draw, (1040, 280), (1160, 310), "#333333")
    draw.line((800, 370, 800, 450), fill="#333333", width=2)
    draw.line((520, 610, 1525, 610), fill="#dddddd", width=1)

    draw.rectangle((520, 650, 1525, 850), fill="#fafafa", outline="#555555", width=2)
    draw.text((545, 674), "章节说明", font=load_font(24), fill="#111111")
    caption_lines = wrap_text(draw, fig.note, load_font(21), 920)
    for idx, line in enumerate(caption_lines[:4]):
        draw.text((545, 718 + idx * 30), line, font=load_font(21), fill="#222222")
    draw.line((70, 905, 1530, 905), fill="#cccccc", width=1)
    draw.text((78, 922), "图示重点：页面操作、业务规则、数据库记录和验证要点保持一致。", font=load_font(20), fill="#444444")
    image.save(REPORT_FIGURE_DIR / fig.filename)
    return REPORT_FIGURE_DIR / fig.filename


def copy_screenshots() -> None:
    for item in SCREENSHOT_FIGURES:
        src = SOURCE_FIGURE_DIR / item.filename
        if src.exists():
            shutil.copy2(src, REPORT_FIGURE_DIR / item.filename)


def strip_existing_database_sections(text: str) -> str:
    text = re.sub(r"\n## 1\.5 数据库设计与核心数据模型.*?(?=\n## 1\.3 |\n# 四、|\Z)", "\n", text, flags=re.S)
    text = re.sub(r"\n## 4\.15 数据库实现与数据一致性验证.*?(?=\n## 4\.16 |\n## 4\.5 |\n# 五、|\Z)", "\n", text, flags=re.S)
    text = re.sub(r"\n## 5\.5 个人负责模块与数据库支撑关系.*?(?=\n## |\n# 五、|\Z)", "\n", text, flags=re.S)
    return text


def cleanup_repetition(text: str) -> str:
    repetitive = (
        "开题阶段的设想主要集中在电商退换货智能客服的可行原型，强调意图识别、知识检索、多轮对话和稳定兜底；"
        "进入最终实现阶段以后，系统已经不再只是一个聊天问答页面，而是围绕顾客、管理员、订单、售后申请、凭证、工单、日志和 AI 辅助建议形成完整闭环。"
        "因此，在说明 "
    )
    text = text.replace("技术架构 是", "技术架构是")
    replacements = {
        f"{repetitive}AI 辅助边界 时，既要从用户能看到的页面讲起，也要把后端服务、数据库表、接口返回和异常兜底讲清楚。":
            "本小节重点说明 AI 在系统中的定位：它用于生成建议、摘要和表达增强，但不越过后端业务规则，也不直接改变售后状态。",
        f"{repetitive}登录注册与权限控制 时，既要从用户能看到的页面讲起，也要把后端服务、数据库表、接口返回和异常兜底讲清楚。":
            "登录注册与权限控制是双端系统的入口，决定顾客和管理员看到的页面、可调用的接口以及可操作的数据范围。",
        f"{repetitive}顾客端我的售后 时，既要从用户能看到的页面讲起，也要把后端服务、数据库表、接口返回和异常兜底讲清楚。":
            "顾客端我的售后承担用户发起售后、补充凭证、查看进度和评价服务的职责，是售后闭环的起点。",
        f"{repetitive}管理员售后审核工作台 时，既要从用户能看到的页面讲起，也要把后端服务、数据库表、接口返回和异常兜底讲清楚。":
            "管理员售后审核工作台承担售后单集中处理职责，是顾客端申请进入平台审核流程后的核心后台页面。",
        f"{repetitive}咨询工作台 时，既要从用户能看到的页面讲起，也要把后端服务、数据库表、接口返回和异常兜底讲清楚。":
            "咨询工作台把聊天会话、订单上下文、知识命中、AI 建议和处理轨迹放在同一业务链路中。",
        f"{repetitive}知识库管理与 RAG 依据展示 时，既要从用户能看到的页面讲起，也要把后端服务、数据库表、接口返回和异常兜底讲清楚。":
            "知识库管理与 RAG 依据展示用于说明客服回复和 AI 草稿的规则来源，避免回答变成黑盒生成。",
        f"{repetitive}图片风险扫描与 C2PA 可信度检测 时，既要从用户能看到的页面讲起，也要把后端服务、数据库表、接口返回和异常兜底讲清楚。":
            "图片风险扫描与 C2PA 可信度检测用于对用户上传凭证做预审提示，但不把预审信号写成绝对鉴伪结论。",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def renumber_headings(text: str, chapter_no: int) -> str:
    lines = text.splitlines()
    h2_count = 0
    h3_counts: dict[int, int] = {}
    current_h2 = ""
    out: list[str] = []
    for line in lines:
        if line.startswith("## "):
            h2_count += 1
            title = re.sub(r"^\d+(?:\.\d+)*\s+", "", line[3:]).strip()
            current_h2 = f"{chapter_no}.{h2_count}"
            h3_counts[h2_count] = 0
            out.append(f"## {current_h2} {title}")
        elif line.startswith("### "):
            title = re.sub(r"^\d+(?:\.\d+)*\s+", "", line[4:]).strip()
            if current_h2:
                current_index = int(current_h2.split(".")[1])
                h3_counts[current_index] += 1
                out.append(f"### {current_h2}.{h3_counts[current_index]} {title}")
            else:
                out.append(line)
        else:
            out.append(line)
    return "\n".join(out) + "\n"


def renumber_table_titles(markdown: str) -> str:
    counter = 0

    def repl(match: re.Match[str]) -> str:
        nonlocal counter
        counter += 1
        return f"**表 {counter} {match.group(1).strip()}**"

    return re.sub(r"\*\*表\s+\d+\s+([^*]+)\*\*", repl, markdown)


def table(title: str, headers: list[str], rows: list[list[str]]) -> str:
    lines = [f"**{title}**", "", "| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in rows)
    return "\n".join(lines)


def database_design_section() -> str:
    return f"""
## 数据库设计与核心数据模型

数据库设计是本系统从客服问答原型升级为复杂售后业务系统的关键。系统没有把订单、售后、凭证、工单和日志简单保存在单个 JSON 字段中，而是围绕真实售后流程拆分为用户、订单、售后申请、处理日志、补充凭证、人工工单、AI 回复草稿、知识文档、检索日志、AI 调用日志、处理轨迹、服务评价和商品质量预警等多类表。这样既能支撑顾客端和管理员端的不同视角，也能在答辩时通过数据库记录证明每个业务动作确实落库、可追踪、可复盘。

在核心业务链路中，user_account 表表示系统操作者，demo_order 表表示顾客订单，after_sale_application 表表示售后申请主对象，after_sale_process_log 表记录每一次状态变化和人工处理动作，after_sale_evidence 表记录顾客补充的图片、物流单号和文字凭证，service_ticket 表承接复杂售后和投诉转人工场景。AI 增强链路则由 knowledge_doc、retrieval_log、ai_call_log、reply_draft 和 process_trace 等表支撑，用于保存知识依据、模型调用结果、回复草稿和处理轨迹。

图 [FIG:fig_b_database_er] 展示了数据库核心实体之间的关系。数据库模型以订单和售后申请为中心：demo_order 提供售后判断所需的商品、支付、物流和订单状态；after_sale_application 保存售后主状态；after_sale_process_log 与 after_sale_evidence 分别保存过程日志和凭证材料；reply_draft、retrieval_log、ai_call_log 与 process_trace 记录 AI 辅助过程。通过这些关系，系统可以从一个售后单追溯到订单、顾客、凭证、工单、知识依据、AI 调用和处理日志，形成完整的业务证据链。

{table("表 3 核心数据库表职责说明", ["表名", "类型", "主要职责", "支撑页面/模块"], [
    ["user_account", "用户与权限表", "保存管理员、顾客账号、角色、账号状态和认证相关信息", "登录注册、权限控制、顾客端、管理员端"],
    ["demo_order", "订单表", "保存订单号、用户、商品、SKU、金额、支付状态、物流状态、售后状态等", "我的售后、订单管理、客服咨询、售后判断"],
    ["after_sale_application", "售后申请主表", "保存售后单号、订单、用户、服务类型、原因、状态、优先级、SLA 截止时间和风险等级", "顾客售后中心、管理员审核台、SLA 队列"],
    ["after_sale_process_log", "售后处理日志表", "记录提交、审核、补证、通过、驳回、完成、采纳草稿、创建工单等动作", "售后时间线、审计记录、答辩验证"],
    ["after_sale_evidence", "售后凭证表", "保存顾客补充的文字、图片、物流单号和其他证据", "补充凭证、证据审核、图片风险"],
    ["service_ticket", "人工工单表", "保存投诉、物流异常、复杂争议等人工处理事项", "人工工单、SLA 协同、管理员跟进"],
    ["reply_draft", "AI 回复草稿表", "保存本地模板或 AI 生成的回复建议、状态、风险等级和知识依据", "AI 副驾驶、审核工作台、日志审计"],
    ["knowledge_doc", "知识文档表", "保存退货、换货、退款、物流、投诉等知识规则", "知识库管理、RAG 检索、回复草稿"],
    ["retrieval_log", "检索日志表", "保存检索词、命中文档、分数、排序和命中原因", "RAG 依据展示、日志诊断"],
    ["ai_call_log", "AI 调用日志表", "保存模型提供方、模型名、请求摘要、响应摘要、耗时、错误和兜底状态", "AI 质检、日志诊断、草稿审计"],
    ["process_trace", "处理轨迹表", "保存聊天或售后处理链路中的上下文解析、意图识别、知识检索、AI 生成等步骤", "咨询工作台、服务日志诊断"],
    ["service_review", "服务评价表", "保存售后完成后的顾客评分和评价内容", "顾客评价、客户画像"],
    ["product_issue_alert", "商品质量预警表", "聚合同一商品的高频问题、风险等级和建议动作", "商品质量预警、运营复盘"],
    ["after_sale_risk_assessment", "售后风险评估表", "保存售后风险分、风险等级、风险原因和建议处理方式", "SLA 队列、风险任务、审核辅助"],
    ["evidence_audit", "凭证审核表", "保存凭证充分性、真实性、AI 生成风险、篡改风险和补证建议", "凭证审核、图片风险、人工复核"],
])}

从表的职责可以看出，系统并不是把所有售后信息集中在单个主表中，而是按照“主对象、过程日志、凭证材料、人工接管、AI 依据、运营复盘”拆分。这样的设计使每类数据都有明确边界：主表负责当前状态，日志表负责历史追踪，凭证表负责材料证据，工单表负责人工协同，AI 相关表负责可解释和可审计，运营类表负责后续复盘。

{table("表 4 售后业务核心表结构字典", ["表名", "字段", "含义", "设计说明"], [
    ["after_sale_application", "id、application_no", "主键和售后单号", "作为售后申请唯一标识，并用于页面展示和客服沟通"],
    ["after_sale_application", "order_id、user_id", "关联订单和申请用户", "关联 demo_order 与 user_account，限制顾客只能查看和操作自己的售后"],
    ["after_sale_application", "service_type、reason_code、reason_text", "服务类型和申请原因", "区分退货、换货、退款、投诉等场景，保存用户填写的具体问题"],
    ["after_sale_application", "status、priority、sla_deadline", "状态、优先级和 SLA 截止时间", "支撑状态机流转、风险任务识别和前端队列展示"],
    ["after_sale_application", "refund_amount、approved_amount", "申请金额和批准金额", "作为管理员审核参考，不由 AI 自动决定"],
    ["after_sale_application", "assigned_to、ticket_id、diagnosis_id", "处理人、工单和诊断关联", "支撑人工接管、前置诊断和审核台上下文"],
    ["after_sale_application", "ai_summary、risk_level、closed_at", "AI 摘要、风险等级和关闭时间", "用于审核辅助、SLA 风险识别和完成状态记录"],
    ["after_sale_process_log", "application_id、operator_id、operator_name、operator_role", "售后单和操作人", "区分顾客、管理员、系统或 AI 动作，便于审计"],
    ["after_sale_process_log", "action、from_status、to_status、remark", "动作、状态变化和备注", "记录提交、通过、驳回、补证、创建工单、使用草稿等动作"],
    ["after_sale_evidence", "application_id、evidence_type、file_url、content、uploaded_by", "凭证归属、类型、文件和内容", "保存图片、视频、文字、物流单号等补充材料"],
])}

上述字段来自当前 sql/schema.sql。与任务书草稿相比，实际代码中售后原因描述字段为 reason_text，处理人字段为 assigned_to，日志状态字段为 from_status 与 to_status，凭证提交人字段为 uploaded_by。报告按真实 SQL 修正这些名称，避免把草稿字段误写成已实现字段。

{table("表 5 AI 与知识检索相关表结构字典", ["表名", "关键字段", "作用", "与业务的关系"], [
    ["knowledge_doc", "title、doc_type、intent_code、question、answer、content、keywords、priority、status", "保存售后规则、FAQ 和客服知识", "为聊天回复和 AI 草稿提供依据"],
    ["retrieval_log", "session_id、message_id、query_text、doc_id、rank_no、score、hit_reason、doc_title_snapshot", "记录每次知识检索结果", "证明回复不是无依据生成"],
    ["ai_call_log", "provider、model_name、request_summary、response_summary、status、latency_ms、error_message", "记录模型调用过程", "支撑 AI 质检和异常回退说明"],
    ["reply_draft", "application_id、ticket_id、draft_content、source_type、knowledge_refs、risk_level、ai_status、status", "保存回复草稿和使用状态", "管理员采纳或废弃后留下审计轨迹"],
    ["process_trace", "session_id、message_id、step_name、step_status、detail_json、created_at", "记录聊天处理步骤", "展示上下文解析、意图识别、知识检索、图片风险和 AI 生成链路"],
])}

AI 相关表的设计重点是可解释和可回退。knowledge_doc 提供规则依据，retrieval_log 保存检索过程，ai_call_log 保存模型调用结果，reply_draft 保存管理员可审计的草稿，process_trace 记录一次咨询或售后处理经过的关键步骤。即使外部模型调用失败，系统仍然能够通过本地模板和日志记录保持主业务可用。

{table("表 6 主外键与业务关联关系表", ["关系", "说明", "业务价值"], [
    ["user_account.id -> demo_order.user_id", "一个顾客可以拥有多个订单", "顾客端只能查看自己的订单"],
    ["demo_order.id -> after_sale_application.order_id", "一个订单可以关联售后申请", "售后判断基于真实订单上下文"],
    ["user_account.id -> after_sale_application.user_id", "一个顾客可以提交多个售后", "防止越权提交或查看他人售后"],
    ["after_sale_application.id -> after_sale_process_log.application_id", "一个售后单对应多条处理日志", "保存完整处理时间线"],
    ["after_sale_application.id -> after_sale_evidence.application_id", "一个售后单可以有多条凭证", "支持补材料和证据链"],
    ["after_sale_application.ticket_id -> service_ticket.id", "当前 SQL 通过售后表 ticket_id 回指工单，未额外声明外键约束", "支持复杂售后人工接管，同时保留需要后续补强的数据库约束点"],
    ["after_sale_application.id -> reply_draft.application_id", "一个售后可生成多条回复草稿", "支持 AI 建议采纳/废弃审计"],
    ["knowledge_doc.id -> retrieval_log.doc_id", "一条知识可被多次命中", "支持知识命中统计和依据展示"],
    ["chat_session.id -> ai_call_log.session_id", "一次会话可包含多次 AI 调用", "追踪模型状态、失败原因和兜底情况"],
    ["after_sale_evidence.id -> evidence_audit.evidence_id", "一条凭证可对应审核结果", "支持图片风险和凭证可信度预审"],
])}

图 [FIG:fig_c_after_sale_state_machine] 进一步展示了售后状态机。售后申请从 SUBMITTED 进入 UNDER_REVIEW 后，管理员可以通过、驳回、要求补证或创建人工工单；顾客补充凭证后可以回到审核流程；通过后由管理员确认完成，完成后顾客才能评价。这个设计避免顾客直接修改通过状态，也避免管理员在没有理由的情况下驳回售后。

{table("表 7 售后状态机与允许操作表", ["当前状态", "顾客可操作", "管理员可操作", "数据库约束说明"], [
    ["SUBMITTED", "查看进度", "开始审核、驳回、要求补证、通过", "新申请必须存在订单和用户关联"],
    ["UNDER_REVIEW", "查看进度", "通过、驳回、要求补证、创建工单、生成草稿", "只能由管理员改变审核状态"],
    ["NEED_MORE_EVIDENCE", "补充凭证、查看补证要求", "查看凭证、创建工单", "补证写入 after_sale_evidence 和 after_sale_process_log"],
    ["APPROVED", "查看处理结果", "确认完成、创建工单", "不允许顾客直接修改通过状态"],
    ["REJECTED", "查看驳回原因", "查看记录", "驳回必须写明原因并留日志"],
    ["COMPLETED", "提交服务评价", "查看评价和复盘", "完成后才允许写入 service_review"],
    ["TICKET_CREATED", "查看人工跟进提示", "更新工单、继续审核", "工单状态变化要回写售后日志"],
])}

{table("表 8 数据库索引与查询场景表", ["查询场景", "涉及表", "当前 SQL 中的索引或全文索引", "说明"], [
    ["顾客查看自己的订单", "demo_order", "idx_order_user(user_id)", "顾客端我的订单按用户过滤"],
    ["顾客查看自己的售后", "after_sale_application", "idx_after_sale_application_user(user_id, status)", "支持按用户和状态查询售后列表"],
    ["管理员按状态筛选售后", "after_sale_application", "idx_after_sale_application_status(status, priority, updated_at)", "支持审核队列分页和排序"],
    ["SLA 查询即将超时任务", "after_sale_application", "idx_after_sale_application_sla(sla_deadline, status)", "支持按截止时间识别风险任务"],
    ["查询售后时间线", "after_sale_process_log", "idx_after_sale_process_log_application(application_id, created_at)", "支持详情页按时间展示日志"],
    ["查询售后凭证", "after_sale_evidence", "idx_after_sale_evidence_application(application_id, created_at)", "支持管理员审核凭证"],
    ["查询人工工单", "service_ticket", "idx_ticket_status(status, priority, created_at)", "支持工单列表筛选和排序"],
    ["知识检索", "knowledge_doc", "idx_doc_intent(intent_code)、idx_doc_status(status, deleted)、ft_doc_search", "支持按意图、状态和全文内容检索规则"],
    ["日志诊断", "ai_call_log / retrieval_log / process_trace", "idx_ai_status、idx_retrieval_message、idx_trace_session(session_id, created_at)", "支持按会话追踪 AI 与检索过程"],
])}

索引设计围绕系统高频查询展开。顾客端高频访问自己的订单和售后，因此需要按 user_id 过滤；管理员端高频按状态、优先级和更新时间查看审核队列；SLA 页面需要按截止时间识别风险任务；详情页需要按 application_id 快速读取处理日志和凭证。当前 sql/schema.sql 已经包含这些核心索引，个别更复杂的组合索引仍可在后续真实数据量增大后继续调优。
""".strip()


def database_consistency_section() -> str:
    return f"""
## 数据库实现与数据一致性验证

数据库实现不是只在报告中列出表名，而是通过 Service 层业务方法把页面动作和多表写入连接起来。当前后端在提交售后、审核通过、驳回、要求补证、补充凭证、创建工单、生成草稿、采纳/废弃草稿、提交评价等方法上使用 @Transactional，说明这些动作并非孤立更新单字段，而是需要保证主记录、日志、凭证、草稿或工单之间的一致性。

售后系统的数据一致性主要体现在“状态变化必须伴随日志记录”。例如，顾客提交售后不仅要写入 after_sale_application，还要写入 after_sale_process_log；管理员要求补材料不仅要修改状态，还要保存补证要求；创建人工工单时，service_ticket 与售后日志必须同步出现。通过这种方式，系统避免出现页面状态变化但无法追溯原因的情况。

{table("表 9 数据一致性与事务控制表", ["业务动作", "涉及表", "一致性要求", "处理方式"], [
    ["顾客提交售后", "after_sale_application、after_sale_process_log、demo_order", "创建售后主记录后必须写入提交日志，并更新订单售后状态", "Service 层统一校验订单归属和状态，并使用事务包裹"],
    ["管理员要求补材料", "after_sale_application、after_sale_process_log", "状态改为 NEED_MORE_EVIDENCE，同时记录补证要求", "后端限制只有审核中售后才能要求补证"],
    ["顾客补充凭证", "after_sale_evidence、after_sale_process_log、after_sale_application", "凭证保存后需要记录补证动作，并允许回到审核状态", "校验售后属于当前顾客"],
    ["管理员驳回售后", "after_sale_application、after_sale_process_log", "驳回状态和驳回原因必须同时保存", "驳回理由不能为空"],
    ["管理员创建工单", "service_ticket、after_sale_application、after_sale_process_log", "工单与售后单必须关联，售后日志需要记录转人工", "工单创建失败时不能只留下售后日志"],
    ["生成 AI 回复草稿", "reply_draft、retrieval_log、ai_call_log", "草稿需要保留知识依据和 AI 状态", "AI 失败时保存本地模板草稿或记录兜底状态"],
    ["采纳/废弃草稿", "reply_draft、after_sale_process_log", "草稿状态变化和管理员动作都要留痕", "采纳不等于自动通过售后"],
    ["提交服务评价", "service_review、after_sale_application", "只有已完成售后才能评价，且评价属于当前顾客", "后端校验状态和用户归属"],
])}

图 [FIG:fig_d_submit_after_sale_sequence] 展示顾客提交售后的接口链路。顾客端页面通过 Axios 发送 POST /customer/after-sales，后端 Controller 获取当前登录用户，Service 层校验订单归属和订单状态，然后写入售后主表、处理日志并更新订单售后状态。这个链路证明顾客端和管理员端使用同一套后端对象和数据库表，只是在权限和视角上不同。

图 [FIG:fig_e_reply_draft_audit_flow] 展示 AI 草稿生成与审计链路。AI 相关表不直接决定售后状态，而是保存知识依据、模型调用结果和回复草稿。管理员采纳草稿只表示采纳一段回复建议，不等于自动通过售后；真正的售后状态仍由审核动作和状态机约束决定。

图 [FIG:fig_f_image_risk_c2pa_flow] 展示图片风险和 C2PA 预审的数据流。图片风险和 C2PA 检测只作为可信度预审信号，不作为绝对鉴伪结论。系统会把风险信号、补证建议和人工复核要求写入凭证审核或处理轨迹，而不是直接用 AI 判断结果替代管理员审核。

图 [FIG:fig_g_sla_ticket_flow] 展示 SLA 与人工工单协同流程。SLA 队列基于状态、优先级、风险评估和截止时间聚合，服务于管理员处理顺序；人工工单用于承接投诉、物流异常和复杂争议，处理后再回到售后审核和日志回写链路。
""".strip()


def personal_database_support_section() -> str:
    return f"""
## 个人负责模块与数据库支撑关系

为了突出个人负责模块不是页面堆叠，本报告把刘剑宇负责的四条主线与数据库表建立对应关系。每个模块都能在数据库中找到主表、日志表或审计表支撑，说明个人工作不仅是界面实现，也包括业务建模、状态约束、AI 边界和可追溯证据链设计。

{table("表 10 数据库对个人负责模块的支撑关系表", ["刘剑宇负责模块", "关键数据库表", "数据库支撑点", "体现的工程能力"], [
    ["真实双端售后业务闭环", "demo_order、after_sale_application、after_sale_process_log、after_sale_evidence、service_review", "支撑顾客提交、管理员审核、补证、完成和评价", "业务建模、状态机、权限边界、日志追踪"],
    ["SLA 跟进与人工工单协同", "after_sale_application、after_sale_risk_assessment、service_ticket、after_sale_process_log", "支撑超时任务、风险队列、人工接管和工单回写", "优先级设计、任务队列、人工协作"],
    ["AI 副驾驶可审计回复草稿", "reply_draft、knowledge_doc、retrieval_log、ai_call_log、process_trace", "支撑知识依据、AI 生成、本地兜底、采纳/废弃审计", "AI 工程化、可解释性、降级能力"],
    ["聊天图片风险扫描与 C2PA 可信度检测", "after_sale_evidence、evidence_audit、process_trace、after_sale_process_log", "支撑图片凭证、风险信号、补证建议和人工复核", "证据链设计、风险提示、可信度边界"],
])}
""".strip()


def split_sections(text: str) -> dict[str, str]:
    matches = list(re.finditer(r"(?m)^# (.+)$", text))
    sections: dict[str, str] = {}
    for idx, match in enumerate(matches):
        heading = match.group(1).strip()
        start = match.end()
        end = matches[idx + 1].start() if idx + 1 < len(matches) else len(text)
        sections[heading] = text[start:end].strip()
    return sections


def build_markdown() -> str:
    raw = SOURCE_MD.read_text(encoding="utf-8")
    raw = strip_existing_database_sections(raw)
    raw = cleanup_repetition(raw)
    sections = split_sections(raw)

    intro = sections.get("一、课题介绍", "")
    team = sections.get("二、团队分工", "")
    design = sections.get("三、系统设计", "")
    impl = sections.get("四、实现情况", "")
    future = sections.get("五、未来与展望", "")

    design = design.replace("## 1.3 前后端请求与权限边界", database_design_section() + "\n\n## 前后端请求与权限边界")
    design = design.replace("## 1.4 AI 辅助边界与本地兜底", "## AI 辅助边界与本地兜底")
    design = design.replace("## 1.2 技术架构设计", "## 技术架构设计")
    design = design.replace("## 1.1 选题背景与结项目标", "## 选题背景与结项目标")

    impl = impl.replace("## 3.1 登录注册与双端权限控制", "## 登录注册与双端权限控制")
    for old in ["## 3.2 ", "## 3.3 ", "## 3.4 ", "## 3.5 ", "## 3.6 ", "## 3.7 ", "## 3.8 ", "## 3.9 ", "## 3.10 ", "## 3.11 ", "## 3.12 ", "## 3.13 ", "## 3.14 ", "## 3.15 "]:
        impl = impl.replace(old, "## ")
    impl = impl.replace("## 4.1 ", "## ")
    impl = impl.replace("## 4.2 ", "## ")
    impl = impl.replace("## 4.3 ", "## ")
    impl = impl.replace("## 4.4 ", "## ")
    impl = impl.replace("## 4.5 测试与验证", database_consistency_section() + "\n\n## 测试与验证")
    impl = impl.replace("## 4.6 项目特色与创新点", "## 项目特色与创新点")
    impl = impl.replace("## 4.14 个人负责模块补充说明", personal_database_support_section() + "\n\n## 个人负责模块补充说明")
    impl = re.sub(r"### \d+\.\d+\.\d+ ", "### ", impl)
    impl = re.sub(r"### \d+\.\d+ ", "### ", impl)

    body = "\n\n".join(
        [
            "# 一、课题介绍",
            renumber_headings(intro, 1).strip(),
            "# 二、团队分工",
            renumber_headings(team, 2).strip(),
            "# 三、系统设计",
            renumber_headings(design, 3).strip(),
            "# 四、实现情况",
            renumber_headings(impl, 4).strip(),
            "# 五、未来与展望",
            renumber_headings(future, 5).strip(),
        ]
    )
    return renumber_table_titles(body.strip()) + "\n"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


class ReportDoc:
    def __init__(self) -> None:
        self.doc = Document()
        self.figure_no = 1
        self.table_no = 1
        self.figure_numbers: dict[str, int] = {}
        self.setup()

    def setup(self) -> None:
        section = self.doc.sections[0]
        section.top_margin = Cm(2.45)
        section.bottom_margin = Cm(2.3)
        section.left_margin = Cm(2.7)
        section.right_margin = Cm(2.35)
        for style_name in ["Normal", "Body Text"]:
            style = self.doc.styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = self.doc.styles[style_name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.color.rgb = RGBColor(0, 0, 0)

    def font(self, run, size: float = 10.5, bold: bool = False, east: str = "宋体") -> None:
        run.font.name = east
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
        run.font.size = Pt(size)
        run.bold = bold

    def add_cover(self) -> None:
        for _ in range(1):
            self.doc.add_paragraph()
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        table.style = "Table Grid"
        cell = table.cell(0, 0)
        set_cell_width(cell, 4.8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        r = cell.paragraphs[0].add_run("成绩：\n")
        self.font(r, size=16)
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("嘉兴学院")
        self.font(r, size=42, east="黑体")
        for _ in range(2):
            self.doc.add_paragraph()
        for line in [
            "课程名称：复杂软件系统实践",
            "任课老师：王向东",
            "题目：基于检索增强的电商退换货智能客服系统设计与实现",
            "",
            "学院：人工智能学院",
            "年级：2025届",
            "专业及班级：软件231",
            "组员姓名学号：刘剑宇（学号提交前补充）",
        ]:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if line:
                r = p.add_run(line)
                self.font(r, size=16)
        for _ in range(3):
            self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("提交时间：2026年5月")
        self.font(r, size=16)
        self.doc.add_section(WD_SECTION.NEW_PAGE)

    def add_toc(self) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("目录")
        self.font(r, size=16, bold=True, east="黑体")
        toc_items = [
            "一、课题介绍",
            "二、团队分工",
            "三、系统设计",
            "四、实现情况",
            "五、未来与展望",
            "评分表",
        ]
        for item in toc_items:
            p = self.doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.55
            r = p.add_run(item)
            self.font(r, size=12)
        self.doc.add_section(WD_SECTION.NEW_PAGE)

    def h1(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        self.font(r, size=15, bold=True, east="黑体")

    def h2(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text)
        self.font(r, size=13, bold=True, east="黑体")

    def h3(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        self.font(r, size=11, bold=True, east="黑体")

    def para(self, text: str) -> None:
        if not text.strip():
            return
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(21)
        p.paragraph_format.line_spacing = 1.22
        p.paragraph_format.space_after = Pt(2)
        for part in re.split(r"(\[FIG:[^\]]+\])", text.strip()):
            if not part:
                continue
            match = re.fullmatch(r"\[FIG:([^\]]+)\]", part)
            if match:
                stem = match.group(1)
                fig_no = self.figure_numbers.get(stem)
                if fig_no:
                    r = p.add_run(f"图 {fig_no}")
                    self.font(r, size=10.5, bold=True)
                else:
                    r = p.add_run(part)
                    self.font(r)
            else:
                r = p.add_run(part)
                self.font(r)

    def add_figure(self, fig: ReportFigure) -> None:
        path = REPORT_FIGURE_DIR / f"{fig.stem}.png"
        if not path.exists():
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Cm(14.8))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(2)
        self.figure_numbers[fig.stem] = self.figure_no
        r = cap.add_run(f"图 {self.figure_no} {fig.title}")
        self.font(r, size=9.5, bold=True)
        self.figure_no += 1
        note = self.doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.paragraph_format.space_after = Pt(6)
        r = note.add_run(fig.note)
        self.font(r, size=9)

    def add_screenshot(self, item: ScreenshotFigure) -> None:
        path = REPORT_FIGURE_DIR / item.filename
        if not path.exists():
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Cm(15.2))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(2)
        r = cap.add_run(f"图 {self.figure_no} {item.title}")
        self.font(r, size=9.5, bold=True)
        self.figure_no += 1
        note = self.doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = note.add_run(item.note)
        self.font(r, size=9)

    def add_chapter_figure(self, item: ChapterFigure) -> None:
        path = REPORT_FIGURE_DIR / item.filename
        if not path.exists():
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Cm(15.25))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(2)
        self.figure_numbers[item.stem] = self.figure_no
        r = cap.add_run(f"图 {self.figure_no} {item.title}")
        self.font(r, size=9.5, bold=True)
        self.figure_no += 1
        note = self.doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.paragraph_format.space_after = Pt(6)
        r = note.add_run(item.note)
        self.font(r, size=9)

    def markdown_table(self, block: str) -> None:
        rows: list[list[str]] = []
        explicit_title: str | None = None
        block = block.strip()
        if block.startswith("**") and "**" in block[2:]:
            first, _, rest = block.partition("\n")
            explicit_title = first.strip("*")
            block = rest
        for raw in block.splitlines():
            stripped = raw.strip()
            if not stripped.startswith("|"):
                continue
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if cells and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
                continue
            rows.append(cells)
        if not rows:
            return
        title = explicit_title
        if title and title.startswith("表 "):
            # Use the explicit table title from markdown and sync the counter.
            match = re.match(r"表\s+(\d+)", title)
            if match:
                self.table_no = max(self.table_no, int(match.group(1)) + 1)
        elif title:
            title = f"表 {self.table_no} {title}"
            self.table_no += 1
        self.add_table(rows[0], rows[1:], title)

    def add_table(self, headers: list[str], rows: list[list[str]], title: str | None = None) -> None:
        if title:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            r = p.add_run(title)
            self.font(r, size=10.5, bold=True)
        col_count = len(headers)
        table = self.doc.add_table(rows=1, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for i, header in enumerate(headers):
            self.set_cell(table.rows[0].cells[i], header, True, "F2F2F2")
        for row in rows:
            cells = table.add_row().cells
            for i in range(col_count):
                self.set_cell(cells[i], row[i] if i < len(row) else "")
        self.doc.add_paragraph()

    def set_cell(self, cell, text: str, bold: bool = False, fill: str | None = None) -> None:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if fill:
            set_cell_shading(cell, fill)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.12
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        size = 8.5 if len(text) > 48 else 9.2
        self.font(r, size=size, bold=bold)

    def add_markdown(self, text: str) -> None:
        blocks = re.split(r"\n{2,}", text)
        for block in blocks:
            line = block.strip()
            if not line:
                continue
            if line.startswith("# "):
                self.h1(line[2:].strip())
            elif line.startswith("## "):
                self.h2(line[3:].strip())
            elif line.startswith("### "):
                self.h3(line[4:].strip())
            elif line.startswith("**") and line.endswith("**") and "\n|" not in line:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True
                r = p.add_run(line.strip("*"))
                self.font(r, size=10.5, bold=True)
            elif line.startswith("[[FIG:") or line.startswith("[[[FIG:"):
                stem = line.replace("[", "").replace("]", "").removeprefix("FIG:").strip()
                fig = next((item for item in ENGINEERING_FIGURES if item.stem == stem), None)
                if fig:
                    self.add_figure(fig)
            elif line.startswith("[[CHAPTER_FIG:"):
                marker, _, rest = line.partition("\n")
                stem = marker.replace("[", "").replace("]", "").removeprefix("CHAPTER_FIG:").strip()
                fig = next((item for item in CHAPTER4_SECTION_FIGURES if item.stem == stem), None)
                if fig:
                    self.add_chapter_figure(fig)
                if rest.strip():
                    self.para(rest.replace("\n", ""))
            elif line.startswith("[[SCREENSHOTS]]"):
                self.h2("真实系统截图补充")
                self.para("以下截图来自本项目运行页面或已有界面图，用于补充说明主要功能模块的实际呈现效果。")
                for item in SCREENSHOT_FIGURES:
                    self.add_screenshot(item)
            elif line.startswith("**") and "\n|" in line:
                self.markdown_table(line)
            elif line.startswith("|"):
                self.markdown_table(line)
            elif line.startswith("- "):
                for item in line.splitlines():
                    self.para(item[2:].strip() if item.startswith("- ") else item.strip())
            else:
                self.para(line.replace("\n", ""))

    def add_score_table(self) -> None:
        self.add_table(
            ["模块", "项目", "要求", "分值", "得分"],
            [
                ["开题报告", "内容", "团队分工合理；方案设计可行", "10", ""],
                ["", "质量", "报告逻辑结构清晰；具有一定创新性", "10", ""],
                ["结项答辩", "团队", "满足设计需求；运行顺畅", "30", ""],
                ["", "", "程序设计规范；UI 设计友好", "10", ""],
                ["", "个人", "对自身开发部分自述内容准确，能够根据设计要点进行全面阐述", "10", ""],
                ["", "", "回答问题思维敏捷、逻辑清晰、语言流畅、正确合理、有自己的见解", "10", ""],
                ["结项报告", "内容", "内容详实、表达正确、逻辑清晰", "10", ""],
                ["", "拓展", "对本项目提出合理、可行建议", "10", ""],
                ["统计", "难度系数", "选题自拟则难度系数为 0.8，否则为 1", "", ""],
                ["", "总分", "各项目累积分值", "100", ""],
            ],
            "评分表",
        )


def add_figure_markers(markdown: str) -> str:
    markdown = re.sub(
        r"(##\s+\d+\.\d+\s+技术架构设计\s*\n)",
        r"\1\n[[FIG:fig_a_system_architecture]]\n\n",
        markdown,
        count=1,
    )
    markdown = markdown.replace("图 [FIG:fig_b_database_er] 展示", "[[FIG:fig_b_database_er]]\n\n图 [FIG:fig_b_database_er] 展示")
    markdown = markdown.replace("图 [FIG:fig_c_after_sale_state_machine] 进一步", "[[FIG:fig_c_after_sale_state_machine]]\n\n图 [FIG:fig_c_after_sale_state_machine] 进一步")
    markdown = markdown.replace("图 [FIG:fig_d_submit_after_sale_sequence] 展示", "[[FIG:fig_d_submit_after_sale_sequence]]\n\n图 [FIG:fig_d_submit_after_sale_sequence] 展示")
    markdown = markdown.replace("图 [FIG:fig_e_reply_draft_audit_flow] 展示", "[[FIG:fig_e_reply_draft_audit_flow]]\n\n图 [FIG:fig_e_reply_draft_audit_flow] 展示")
    markdown = markdown.replace("图 [FIG:fig_f_image_risk_c2pa_flow] 展示", "[[FIG:fig_f_image_risk_c2pa_flow]]\n\n图 [FIG:fig_f_image_risk_c2pa_flow] 展示")
    markdown = markdown.replace("图 [FIG:fig_g_sla_ticket_flow] 展示", "[[FIG:fig_g_sla_ticket_flow]]\n\n图 [FIG:fig_g_sla_ticket_flow] 展示")
    for fig in CHAPTER4_SECTION_FIGURES:
        markdown = re.sub(
            rf"(?m)^(##\s+{re.escape(fig.heading)}\s*)$",
            rf"\1\n\n[[CHAPTER_FIG:{fig.stem}]]",
            markdown,
            count=1,
        )
    markdown = markdown.replace("# 五、未来与展望", "[[SCREENSHOTS]]\n\n# 五、未来与展望")
    markdown = markdown.replace("[[[FIG:", "[[FIG:").replace("]]]", "]]")
    markdown = markdown.replace("图 [FIG:", "[FIG:")
    return markdown


def make_docx(markdown: str) -> None:
    report = ReportDoc()
    report.add_cover()
    report.add_toc()
    report.add_markdown(markdown)
    report.add_score_table()
    report.doc.save(OUT_DOCX)


def stats(markdown: str) -> dict[str, int]:
    return {
        "chinese_chars": sum(1 for char in markdown if "\u4e00" <= char <= "\u9fff"),
        "tables": len(re.findall(r"(?m)^\*\*表\s+\d+", markdown)),
        "engineering_figures": len(ENGINEERING_FIGURES),
        "chapter4_section_figures": sum(1 for item in CHAPTER4_SECTION_FIGURES if (REPORT_FIGURE_DIR / item.filename).exists()),
        "chapter4_section_markers": len(re.findall(r"\[\[CHAPTER_FIG:", markdown)),
        "screenshots": sum(1 for item in SCREENSHOT_FIGURES if (REPORT_FIGURE_DIR / item.filename).exists()),
    }


def render_pdf_and_pages() -> tuple[Path | None, int]:
    qa_dir = ROOT / "tmp" / "final-report-enhanced-render"
    qa_dir.mkdir(parents=True, exist_ok=True)
    for old_page in qa_dir.glob("page-*.png"):
        old_page.unlink()
    old_pdf = qa_dir / (OUT_DOCX.stem + ".pdf")
    if old_pdf.exists():
        old_pdf.unlink()
    soffice = shutil.which("soffice") or r"C:\Program Files\LibreOffice\program\soffice.com"
    if not Path(soffice).exists():
        return None, 0
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(qa_dir), str(OUT_DOCX)],
        check=True,
        cwd=ROOT,
    )
    pdf = qa_dir / (OUT_DOCX.stem + ".pdf")
    pdftoppm = shutil.which("pdftoppm") or r"D:\texlive\2025\bin\windows\pdftoppm.exe"
    if Path(pdftoppm).exists() and pdf.exists():
        subprocess.run([pdftoppm, "-png", "-r", "120", str(pdf), str(qa_dir / "page")], check=True, cwd=ROOT)
    pages = len(list(qa_dir.glob("page-*.png")))
    return pdf if pdf.exists() else None, pages


def write_summary(markdown: str, pdf: Path | None, page_count: int) -> None:
    info = stats(markdown)
    SUMMARY_PATH.write_text(
        "\n".join(
            [
                "# 结项报告图表与数据库增强摘要",
                "",
                f"- 输出 DOCX：{OUT_DOCX}",
                f"- 输出 Markdown：{OUT_MD}",
                f"- 工程图源目录：{REPORT_FIGURE_DIR}",
                f"- 中文字符数：{info['chinese_chars']}",
                f"- 新增工程图：{info['engineering_figures']} 张",
                f"- 第四章章节示意图：{info['chapter4_section_figures']} 张",
                f"- 保留真实系统截图：{info['screenshots']} 张",
                f"- 新增/整理表格：{info['tables']} 张",
                f"- PDF 渲染页数：{page_count}",
                f"- PDF 文件：{pdf if pdf else '未生成'}",
                "",
                "## 新增工程图",
                *[f"- {fig.title}：{fig.stem}.mmd / {fig.stem}.png" for fig in ENGINEERING_FIGURES],
                "",
                "## 第四章章节示意图",
                *[f"- {fig.heading}：{fig.filename}" for fig in CHAPTER4_SECTION_FIGURES],
                "",
                "## 真实性说明",
                "- 数据库字段按 sql/schema.sql 校对：reason_text、assigned_to、from_status/to_status、uploaded_by 等字段采用真实名称。",
                "- service_ticket 当前 SQL 未声明 application_id 外键；报告改写为 after_sale_application.ticket_id 的业务回指关系。",
                "- AI 草稿没有写成自动业务决策，图片风险与 C2PA 只写成可信度预审信号。",
                "- 第四章章节示意图围绕本项目 Spring Boot + Vue 3 + MySQL + LangChain4j、售后、工单、知识库、日志和图片风险预审等实际模块组织。",
            ]
        ),
        encoding="utf-8",
    )
    QUALITY_PATH.write_text(
        "\n".join(
            [
                "# 增强版质量检查",
                "",
                f"- 中文字符数：{info['chinese_chars']}",
                f"- 工程图：{info['engineering_figures']}",
                f"- 第四章章节示意图：{info['chapter4_section_figures']}",
                f"- 第四章章节图标记：{info['chapter4_section_markers']}",
                f"- 真实截图：{info['screenshots']}",
                f"- 表格：{info['tables']}",
                f"- 渲染页数：{page_count}",
                "",
                "- [x] 第四章 4.1 至 4.24 每个小节均插入章节示意图",
                "- [x] 新增数据库核心 ER 图",
                "- [x] 新增售后状态机图",
                "- [x] 新增顾客提交售后时序图",
                "- [x] 新增 AI 草稿生成与审计链路图",
                "- [x] 新增图片风险与 C2PA 数据流图",
                "- [x] 新增 SLA 与人工工单协同流程图",
                "- [x] 新增核心表职责、字段字典、主外键关系、状态机、索引、事务一致性和个人模块支撑表",
                "- [x] 修正章节编号，压缩重复套话",
                "- [x] 未伪造性能、上线或真实业务数据",
                "- [x] 未把 AI 写成最终业务决策者",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    REPORT_FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    for fig in ENGINEERING_FIGURES:
        render_figure(fig)
    for fig in CHAPTER4_SECTION_FIGURES:
        render_chapter_figure(fig)
    copy_screenshots()
    markdown = add_figure_markers(build_markdown())
    OUT_MD.write_text(markdown, encoding="utf-8")
    make_docx(markdown)
    pdf, page_count = render_pdf_and_pages()
    write_summary(markdown, pdf, page_count)
    info = stats(markdown)
    print(f"DOCX: {OUT_DOCX}")
    print(f"Markdown: {OUT_MD}")
    print(f"Figures: {info['engineering_figures']} engineering + {info['chapter4_section_figures']} chapter4 section figures + {info['screenshots']} screenshots")
    print(f"Tables: {info['tables']}")
    print(f"Chinese chars: {info['chinese_chars']}")
    print(f"Rendered pages: {page_count}")
    if info["engineering_figures"] < 7:
        raise SystemExit("Expected at least 7 engineering figures")
    if info["chapter4_section_figures"] != len(CHAPTER4_SECTION_FIGURES):
        raise SystemExit("Expected all chapter 4 section figures to exist")
    if info["chapter4_section_markers"] != len(CHAPTER4_SECTION_FIGURES):
        raise SystemExit("Expected one section figure marker for every chapter 4 subsection")
    if info["tables"] < 8:
        raise SystemExit("Expected at least 8 tables")
    if page_count <= 0:
        raise SystemExit("DOCX render failed")


if __name__ == "__main__":
    main()
