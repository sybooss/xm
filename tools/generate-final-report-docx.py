from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "final-report.docx"


def set_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_para(doc, text="", style=None, align=None, size=10.5, bold=False, first_line=True):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if first_line and style is None and text:
        p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.2
    for line_no, line in enumerate(str(text).split("\n")):
        if line_no:
            p.add_run().add_break()
        run = p.add_run(line)
        set_font(run, size=9.5, bold=bold)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True)
    return p


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目录")
    set_font(run, size=16, bold=True)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "右键更新域可生成目录"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def main():
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)

    add_para(doc, "人工智能学院", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, first_line=False)
    doc.add_paragraph()
    add_para(doc, "复杂软件系统实践", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, first_line=False)
    add_para(doc, "结项报告", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, first_line=False)
    doc.add_paragraph()
    add_para(doc, "项目名称：电商退换货智能客服系统", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, first_line=False)
    add_para(doc, "技术栈：Spring Boot + Vue 3 + MySQL + LangChain4j", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, first_line=False)
    doc.add_paragraph()
    for item in ["学院：人工智能学院", "专业班级：软件231", "成员姓名/学号：", "指导教师：", "提交时间：2026年5月"]:
        add_para(doc, item, size=12, first_line=False)
    doc.add_section(WD_SECTION.NEW_PAGE)

    add_toc(doc)
    doc.add_section(WD_SECTION.NEW_PAGE)

    add_heading(doc, "一、系统简介", 1)
    add_para(doc, "本项目为“电商退换货智能客服系统”，面向电商售后中高频出现的退货申请、换货申请、退款进度、物流异常、规则咨询、投诉与人工转接等场景。系统采用 Spring Boot + Vue 3 + MySQL + LangChain4j 技术栈，目标不是只实现一个聊天框，而是构建一条可演示、可追踪、可维护的售后客服业务闭环。")
    add_para(doc, "系统主流程为：用户登录/注册，绑定订单并发起售后咨询，后端识别售后意图，读取订单和售后上下文，检索知识库规则，生成本地业务判断，调用 LangChain4j 进行 AI 增强，失败时走本地规则兜底，必要时创建人工客服工单，并保存消息、意图、检索、AI 调用和处理轨迹日志。")

    add_heading(doc, "二、主要功能", 1)
    add_heading(doc, "2.1 登录注册与权限控制", 2)
    add_para(doc, "系统提供管理员和客户两类角色。管理员可访问答辩展示中心、系统总览、知识库、订单管理、人工工单、日志中心和 AI 测试；客户注册后默认角色为 CUSTOMER，只能访问咨询工作台和自己的订单售后入口。权限由后端 AuthInterceptor、AuthServiceImpl 和 @OperatorAnno 配合完成，前端路由通过 adminOnly 元信息限制后台页面。")
    add_heading(doc, "2.2 售后咨询工作台", 2)
    add_para(doc, "咨询工作台支持会话创建、订单绑定、消息发送、流式处理进度、AI 增强回复和本地兜底回复。右侧处理洞察面板展示意图识别、上下文承接、订单上下文、知识命中、人工转接和回答过程。")
    add_heading(doc, "2.3 知识库和 RAG 依据展示", 2)
    add_para(doc, "知识库包含退货规则、换货规则、退款规则、物流异常、投诉与人工转接和通用 FAQ 等分类。聊天链路中，后端按用户问题和意图检索知识文档，并将命中依据写入 retrieval_log。前端在咨询工作台和日志中心展示检索依据。")
    add_heading(doc, "2.4 LangChain4j AI 增强与本地兜底", 2)
    add_para(doc, "后端通过 AiServiceImpl 接入 LangChain4j OpenAI-compatible 模型。AI 只作为增强层，不替代业务规则。订单是否存在、是否可退换、是否已有售后申请、是否需要人工工单，均由 Spring Boot 业务服务判断。模型不可用时系统使用本地规则兜底，保证现场稳定演示。")
    add_heading(doc, "2.5 智能工单升级", 2)
    add_para(doc, "当用户表达投诉、人工客服、平台介入、商家一直不处理等诉求时，系统会创建或复用人工客服工单。工单保存会话、订单、意图、优先级、状态、用户问题、AI 摘要和处理建议，使系统从问答工具升级为售后服务闭环。")
    add_heading(doc, "2.6 日志诊断与答辩展示中心", 2)
    add_para(doc, "系统保存 AI 调用日志、知识检索日志和处理轨迹。日志诊断中心聚合展示 AI 成功率、平均耗时、知识命中数量、去重命中文档、平均检索分数、会话轨迹步骤和高频命中文档，同时保留原始日志表格和轨迹时间线，方便答辩时证明系统可解释、可调试。新增 /showcase 答辩展示中心，管理员登录后默认进入该页面，可集中展示系统主题、演示顺序、8 个已实现亮点、系统状态、兜底策略和关键页面入口。")

    add_heading(doc, "三、系统设计与关键实现", 1)
    add_heading(doc, "3.1 后端分层设计", 2)
    add_para(doc, "后端采用 Controller、Service、Mapper、Pojo 分层结构。Controller 负责参数接收和统一返回，Service 负责业务编排和事务边界，Mapper 负责 MyBatis 数据访问，Pojo 保存实体、请求对象、分页对象和统一响应。接口采用资源名词和标准 HTTP 方法，例如 GET /service-tickets、POST /chat-sessions/{id}/messages。")
    add_heading(doc, "3.2 聊天主链路", 2)
    add_para(doc, "聊天主链路集中在 ChatServiceImpl.sendMessage。该方法完成消息保存、上下文解析、意图识别、订单查询、知识库检索、本地回复、AI 增强、工单判定、日志记录和最终响应组装。")
    add_table(doc, ["步骤", "说明"], [
        ["CONTEXT_RESOLVE", "解析多轮上下文，判断是否为追问"],
        ["INTENT_RECOGNIZE", "识别售后意图，如退货、退款、物流、投诉"],
        ["ORDER_CONTEXT", "读取订单状态、物流状态和售后状态"],
        ["KNOWLEDGE_RETRIEVAL", "检索知识库并记录命中依据"],
        ["BUSINESS_TOOL_CALLS", "封装订单查询、知识检索和工单工具结果"],
        ["AI_GENERATION", "调用 LangChain4j 生成增强回复，失败时兜底"],
        ["HUMAN_TICKET_CHECK", "判断是否需要人工工单"],
        ["FINAL_REPLY", "保存助手消息并返回前端"],
    ])
    add_heading(doc, "3.3 数据库设计", 2)
    add_table(doc, ["表名", "用途"], [
        ["user_account", "管理员、客户和注册用户"],
        ["demo_order", "演示订单和订单上下文"],
        ["after_sale_record", "退货、换货、退款、投诉等售后记录"],
        ["knowledge_category / knowledge_doc", "知识分类和规则文档"],
        ["chat_session / chat_message", "会话和消息记录"],
        ["intent_record", "意图识别结果"],
        ["retrieval_log", "知识库检索日志"],
        ["ai_call_log", "AI 调用日志"],
        ["process_trace", "聊天处理轨迹"],
        ["service_ticket", "人工客服工单"],
    ])
    add_heading(doc, "3.4 前端设计", 2)
    add_para(doc, "前端采用 Vue 3、Vite、Pinia、Vue Router、Axios 和 Element Plus。页面包括登录注册、答辩展示中心、咨询工作台、系统总览、知识库、订单管理、人工工单、日志诊断中心和 AI 测试。答辩展示中心和日志诊断中心均采用白底、轻边框、留白和清晰层级，强调克制、精致和产品感。")

    add_heading(doc, "四、测试与验证", 1)
    add_para(doc, "本项目不伪造测试结果。最近一次验证结果如下：")
    add_table(doc, ["验证项", "命令", "结果"], [
        ["后端打包", "cd server; mvn -q -DskipTests package", "通过"],
        ["前端构建", "cd web; npm.cmd run build", "通过；存在 Vite chunk size 提示但不影响构建"],
        ["全链路接口烟测", "powershell -NoProfile -ExecutionPolicy Bypass -File .\\tools\\full-smoke-test.ps1", "FAILED_COUNT=0"],
        ["浏览器主流程烟测", "cd web; npm.cmd run test:browser", "FAILED_COUNT=0"],
        ["浏览器角色权限烟测", "cd web; npm.cmd run test:browser:roles", "FAILED_COUNT=0"],
    ])
    add_para(doc, "全链路接口烟测覆盖登录、注册、系统状态、AI 测试、知识库、订单售后、聊天、多轮追问、自动工单、日志和测试数据清理。浏览器烟测覆盖展示中心、咨询工作台、AI 测试、知识库、订单、工单和日志诊断中心。")

    add_heading(doc, "五、项目特色与创新点", 1)
    highlights = [
        "系统不是单一聊天框，而是完整售后客服业务闭环。",
        "RAG 知识库命中结果可在页面和日志中追溯。",
        "LangChain4j 工具调用把订单查询、知识检索和工单能力接入 AI 上下文。",
        "多轮追问既能通过上下文继承理解省略问题，也能在本地规则下稳定运行。",
        "AI 失败不影响核心业务，系统保留本地规则兜底，适合现场答辩。",
        "投诉和异常物流可以自动升级人工工单，体现真实客服协同。",
        "日志诊断中心提供 AI 调用、检索和处理轨迹，并聚合 AI 成功率、平均耗时、知识命中和会话步骤，便于证明系统可解释、可调试。",
        "答辩展示中心把完整项目能力整理成可讲、可点、可验证的高分入口。",
    ]
    for item in highlights:
        add_para(doc, item)

    add_heading(doc, "六、团队分工", 1)
    add_table(doc, ["成员方向", "主要工作"], [
        ["后端业务与数据库", "设计 MySQL 表结构，实现订单、售后、知识库、会话、工单和日志接口"],
        ["AI 与客服链路", "实现意图识别、多轮上下文、RAG 检索、LangChain4j 接入和本地兜底"],
        ["前端页面与交互", "实现登录注册、咨询工作台、展示中心、知识库、工单、日志诊断和 AI 测试页面"],
        ["测试与文档", "编写 README、接口文档、数据库文档、演示脚本、测试用例和烟测脚本"],
    ])

    add_heading(doc, "七、存在不足与未来拓展", 1)
    futures = [
        "知识库检索目前主要基于 MySQL 文本检索和关键词匹配，后续可加入向量检索。",
        "登录权限目前是轻量自实现，后续可接入 Spring Security、刷新 token 和更细粒度 RBAC。",
        "日志诊断中心已能展示 AI 成功率、平均耗时、知识命中和处理轨迹摘要，后续可继续扩展意图分布、工单趋势、客服处理时长和按日期聚合的趋势看板。",
        "当前演示订单是固定种子数据，后续可接入真实电商订单系统或模拟更复杂的售后状态流转。",
        "现有 AI 工具调用主要将业务工具结果注入 prompt，后续可扩展为更标准的 agent 工具执行链。",
        "前端后台页面还可以继续提升产品质感和数据可视化表现。",
    ]
    for item in futures:
        add_para(doc, item)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
